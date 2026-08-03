#!/usr/bin/env python3
"""Verification for the ToR builder — asserts on the WRITTEN file, not on calls.

Six ToR versions (v0.4 to v0.9) shipped unbranded while every generator run
reported success. That is the failure this test exists to make impossible: a
save() that returns without raising says nothing about whether the logo, the
header or the footer strip are in the file. So every assertion below reopens the
saved .docx and inspects the actual package parts.

The slim-header checks look at both the <mc:Choice> drawing and its
<mc:Fallback> VML twin, because Word stores the contact block twice and removing
only the first leaves the text rendering on the page.

Run: python scripts/test_tor_docx.py
"""

from __future__ import annotations

import os
import re
import sys
import tempfile
import zipfile
from pathlib import Path

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import tor_docx  # noqa: E402

MC_NS = "http://schemas.openxmlformats.org/markup-compatibility/2006"

# Strings that exist only in the letterhead's header contact block.
CONTACT_MARKERS = ["Rue Royale", "0840.619.519", "info@ippf.org", "europe.ippf.org"]

results: list[tuple[bool, str]] = []


def check(condition: bool, label: str) -> None:
    results.append((bool(condition), label))


def build(header: str, out_dir: Path) -> Path:
    builder = tor_docx.TorBuilder(header=header)
    builder.title("Terms of Reference — verification build", f"header={header}")
    builder.h1("At a glance")
    builder.kv_table([("Assignment", "Verification"), ("Budget", "EUR 0")])
    builder.h1("Scope")
    builder.bullets(["First bullet", "Second bullet"])
    builder.numbered(["Step one: do a thing", "Step two: do another"])
    builder.table(["Deliverable", "Due"], [["Inception report", "Week 2"],
                                           ["Final report", "Week 8"]])
    builder.para("Closing paragraph.")
    return builder.save(out_dir / f"tor-{header}.docx")


def parts(path: Path) -> dict[str, bytes]:
    with zipfile.ZipFile(path) as archive:
        return {name: archive.read(name) for name in archive.namelist()}


def header_xml(blobs: dict[str, bytes]) -> str:
    return "\n".join(
        blob.decode("utf8", "replace")
        for name, blob in blobs.items()
        if re.match(r"word/header\d*\.xml$", name)
    )


def main() -> int:
    tmp = Path(tempfile.mkdtemp(prefix="tor-verify-"))

    # --- full header -------------------------------------------------------
    full_path = build("full", tmp)
    check(full_path.exists() and full_path.stat().st_size > 0, "full: file written")
    blobs = parts(full_path)

    check(any(re.match(r"word/header\d*\.xml$", n) for n in blobs),
          "full: header part present in the written package")
    check(any(re.match(r"word/footer\d*\.xml$", n) for n in blobs),
          "full: footer part present in the written package")
    check(any(n.startswith("word/media/") for n in blobs),
          "full: media (logo/footer strip) embedded in the written package")

    header_rels = [n for n in blobs if re.match(r"word/_rels/header\d*\.xml\.rels$", n)]
    check(bool(header_rels), "full: header relationship part present")
    rels_xml = "\n".join(blobs[n].decode("utf8", "replace") for n in header_rels)
    check("image" in rels_xml, "full: logo image relationship resolves from the header")

    full_header = header_xml(blobs)
    check(all(marker in full_header for marker in CONTACT_MARKERS),
          "full: contact block retained (address, company number, email, website)")

    # --- slim header -------------------------------------------------------
    slim_path = build("slim", tmp)
    check(slim_path.exists() and slim_path.stat().st_size > 0, "slim: file written")
    slim_blobs = parts(slim_path)
    slim_header = header_xml(slim_blobs)

    for marker in CONTACT_MARKERS:
        check(marker not in slim_header, f"slim: contact marker {marker!r} gone from ALL header parts")

    # The twin check: neither the modern drawing nor the VML fallback may survive.
    remaining = slim_header.count(f"{{{MC_NS}}}AlternateContent") + slim_header.count("mc:AlternateContent")
    check(remaining == 0, "slim: no mc:AlternateContent left (Choice AND Fallback both removed)")
    check("v:textbox" not in slim_header,
          "slim: VML fallback textbox removed, not just the DrawingML choice")

    # Branding must survive the strip — this is the regression that matters.
    check(any(n.startswith("word/media/") for n in slim_blobs),
          "slim: media still embedded after stripping the contact block")
    slim_rels = "\n".join(
        slim_blobs[n].decode("utf8", "replace")
        for n in slim_blobs if re.match(r"word/_rels/header\d*\.xml\.rels$", n)
    )
    check("image" in slim_rels, "slim: logo relationship survives the contact-block removal")
    check(any(re.match(r"word/footer\d*\.xml$", n) for n in slim_blobs),
          "slim: footer strip part survives")

    # --- content actually landed ------------------------------------------
    from docx import Document
    doc = Document(str(slim_path))
    text = "\n".join(p.text for p in doc.paragraphs)
    check("Terms of Reference" in text, "slim: title text written into the body")
    check("•" in text, "slim: bullet glyph rendered without the missing List Bullet style")
    check(re.search(r"1\.\s*\tStep one", text) or "1." in text,
          "slim: numbered list rendered without the missing List Number style")
    check(len(doc.tables) == 2, "slim: both tables written")
    borders = doc.tables[0]._tbl.tblPr.findall(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblBorders")
    check(bool(borders), "slim: table borders applied directly, not via Table Grid style")

    # --- guard rails -------------------------------------------------------
    try:
        tor_docx.TorBuilder(header="nonsense")
        check(False, "invalid header value rejected")
    except ValueError:
        check(True, "invalid header value rejected")

    passed = sum(1 for ok, _ in results if ok)
    for ok, label in results:
        print(f"  {'PASS' if ok else 'FAIL'}  {label}")
    print(f"\n{passed}/{len(results)} checks passed")
    print(f"artefacts: {tmp}")
    return 0 if passed == len(results) else 1


if __name__ == "__main__":
    sys.exit(main())
