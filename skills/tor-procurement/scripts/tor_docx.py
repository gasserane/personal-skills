"""ToR docx builder scaffold — IPPF Visual Identity 2025.

Proven pattern from the AI-for-Research ToR v0.4-v0.7 generators (July 2026):
each ToR lives as a self-contained generator script next to its deliverable
(<project>/generators/gen_tor_vNN_docx.py), so the docx is always regenerable
from source and never dual-edited.

Usage by the tor-procurement skill (build mode):
1. Copy this file into the project's generators/ folder as gen_tor_v01_docx.py.
2. Keep the constants, helpers and TorBuilder unchanged.
3. Replace build_tor() with the real ToR body as linear builder calls,
   following the canonical section order in references/open-market-checklist.md.
4. Run it; the docx lands one level above the generators/ folder.

Builds on the IPPF EN letterhead, not on a blank Document(). This matters more
than it looks. Fonts and colours can be set in code, but the logo, the header
contact block and the pillar footer strip exist ONLY inside the letterhead asset
— nothing in python-docx can conjure them. Starting from Document() therefore
produced correctly-coloured but unbranded files, and six ToR versions (v0.4 to
v0.9) shipped that way before anyone noticed.

Brand values are imported from ane_package.reporting.brand, never mirrored.
Mirroring a brand constant in a builder is the regression the project CLAUDE.md
names as grounds for QA rejection: the copy cannot know when the template moves.
Importing costs nothing here, because opening the letterhead already requires
ane_package on the path — one dependency, not two.
"""

import argparse
import copy
import os
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


def _bootstrap_ane_package() -> None:
    """Put the work folder on sys.path so ane_package imports from anywhere.

    ane_package is not pip-installed; it resolves only when the work folder is
    the current directory. A generator lives in its own project's generators/
    folder, so a bare import would fail there — which is what the old "keep it
    self-contained" comment was really defending against.

    Copying the brand values was the wrong answer to that problem: the letterhead
    asset lives inside ane_package too, so build mode has to find the work folder
    either way. Finding it once fixes both.

    Resolution order follows the project's own convention (CLAUDE.md path
    constants, ane_package/config.py): the WORK_FOLDER_ROOT environment variable
    first, then a walk up from this file, so a generator copied deep into a
    project tree still resolves without configuration.
    """
    try:
        import ane_package  # noqa: F401
        return
    except ModuleNotFoundError:
        pass

    candidates = []
    env_root = os.environ.get("WORK_FOLDER_ROOT")
    if env_root:
        candidates.append(Path(env_root))
    candidates.extend(Path(__file__).resolve().parents)

    for candidate in candidates:
        if (candidate / "ane_package" / "reporting" / "brand.py").is_file():
            sys.path.insert(0, str(candidate))
            return

    raise ModuleNotFoundError(
        "ane_package not found. Build mode needs it for the IPPF brand template "
        "and the letterhead asset. Set WORK_FOLDER_ROOT to the work folder, or "
        "run this generator from inside it."
    )


_bootstrap_ane_package()

from ane_package.reporting.brand import IPPF_FORMAT_TEMPLATE, brand_asset  # noqa: E402

# Single source of truth — resolved from the template, never restated.
FONT = IPPF_FORMAT_TEMPLATE["fonts"]["default"]["name"]
DREAM_HEX = IPPF_FORMAT_TEMPLATE["colours"]["dream"].lstrip("#")
FIRE_RED_HEX = IPPF_FORMAT_TEMPLATE["colours"]["fire_red"].lstrip("#")
DREAM = RGBColor.from_string(DREAM_HEX)
FIRE_RED = RGBColor.from_string(FIRE_RED_HEX)

# The letterhead ships no "Table Grid", "List Bullet" or "List Number" style, so
# every list glyph and table border below is applied directly. Referring to those
# styles by name raises KeyError against this base.
MC_NS = "http://schemas.openxmlformats.org/markup-compatibility/2006"

OUT_NAME = "2 ToR - [assignment] - v0.1 - DRAFT.docx"


def _set_borders(table, hex_colour: str = "BFBFBF", size: int = 4) -> None:
    """Draw grid lines directly on the table.

    The letterhead carries no "Table Grid" style, so `table.style = "Table Grid"`
    raises KeyError against it. Writing <w:tblBorders> onto the table properties
    gives the same result without depending on a style that is not there.
    """
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(size))
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), hex_colour)
        borders.append(element)
    table._tbl.tblPr.append(borders)


def _shade(cell, hex_colour: str) -> None:
    """Fill a table cell (w:shd) with a solid colour."""
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_colour)
    cell._tc.get_or_add_tcPr().append(shd)


def _cell_text(cell, text: str, bold: bool = False, white: bool = False,
               size: int = 10) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    if white:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def _strip_contact_block(document) -> int:
    """Remove the header contact block for `--header slim`. Returns blocks removed.

    Word stores this content TWICE: an <mc:Choice> holding a modern DrawingML
    text box, and an <mc:Fallback> holding a VML twin for older readers. Removing
    only the first leaves the text rendering, which is how a contact block once
    survived a removal that looked successful. Deleting the whole
    <mc:AlternateContent> parent takes both twins together and cannot desync.

    Safe against the logo: in this letterhead every image sits OUTSIDE the
    AlternateContent blocks, and each block is pure contact text. The function
    verifies that per block anyway rather than trusting the layout, so a future
    letterhead that nests the logo differently loses nothing silently.
    """
    removed = 0
    for section in document.sections:
        parts = [section.header, section.footer,
                 section.first_page_header, section.first_page_footer,
                 section.even_page_header, section.even_page_footer]
        for part in parts:
            if part is None:
                continue
            element = part._element
            for block in element.findall(f".//{{{MC_NS}}}AlternateContent"):
                # Never drop a block that carries an image.
                if block.findall(f".//{qn('a:blip')}"):
                    continue
                block.getparent().remove(block)
                removed += 1
    return removed


def open_letterhead(header: str = "full"):
    """Open the IPPF EN letterhead, clear the body, keep the section properties.

    The final <w:sectPr> holds the header/footer references, margins and page
    setup, so it is the one body child that must survive the clear — remove it
    and the document loses the branding it was opened for.
    """
    document = Document(str(brand_asset("word_letterhead_base")))
    body = document.element.body
    for child in list(body):
        if child.tag in (qn("w:p"), qn("w:tbl")):
            body.remove(child)
    if header == "slim":
        _strip_contact_block(document)
    return document


class TorBuilder:
    """Linear ToR document builder. One method call per block, top to bottom."""

    def __init__(self, header: str = "full"):
        """header: 'full' keeps the contact block; 'slim' is band plus logo.

        A multi-page ToR wants slim — repeating the Brussels address and company
        number on every page eats the top of each page and tells the reader
        nothing new after page 1.
        """
        if header not in ("full", "slim"):
            raise ValueError(f"header must be 'full' or 'slim', got {header!r}")
        self.header_mode = header
        self.doc = open_letterhead(header)
        for section in self.doc.sections:
            section.top_margin = Cm(2.0)
            section.bottom_margin = Cm(2.0)
            section.left_margin = Cm(2.2)
            section.right_margin = Cm(2.2)
        style = self.doc.styles["Normal"]
        style.font.name = FONT
        style.font.size = Pt(10.5)

    def title(self, text: str, subtitle: str = "") -> None:
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = FONT
        run.font.size = Pt(20)
        run.font.bold = True
        run.font.color.rgb = DREAM
        if subtitle:
            p2 = self.doc.add_paragraph()
            r2 = p2.add_run(subtitle)
            r2.font.name = FONT
            r2.font.size = Pt(12)
            r2.font.color.rgb = FIRE_RED

    def h1(self, text: str) -> None:
        p = self.doc.add_paragraph()
        p.space_before = Pt(14)
        run = p.add_run(text)
        run.font.name = FONT
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = DREAM

    def h2(self, text: str) -> None:
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = FONT
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = FIRE_RED

    def para(self, text: str) -> None:
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = FONT
        run.font.size = Pt(10.5)

    def lead(self, label: str, text: str) -> None:
        """Paragraph with a bold run-in lead: 'Label. Rest of sentence.'"""
        p = self.doc.add_paragraph()
        r1 = p.add_run(f"{label} ")
        r1.font.name = FONT
        r1.font.size = Pt(10.5)
        r1.font.bold = True
        r2 = p.add_run(text)
        r2.font.name = FONT
        r2.font.size = Pt(10.5)

    def _list_paragraph(self, glyph: str):
        """Indented paragraph carrying its own glyph.

        The letterhead has no "List Bullet"/"List Number" style. Do not reach for
        them and catch the KeyError: add_paragraph(text, style=...) is not atomic,
        so python-docx appends the paragraph first and assigns the style second,
        leaving an orphan paragraph behind when the style is missing. Writing the
        glyph directly means exactly one paragraph is ever added.
        """
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.first_line_indent = Cm(-0.4)
        run = p.add_run(glyph)
        run.font.name = FONT
        run.font.size = Pt(10.5)
        return p

    def bullets(self, items: list[str]) -> None:
        for item in items:
            p = self._list_paragraph("•\t")
            run = p.add_run(item)
            run.font.name = FONT
            run.font.size = Pt(10.5)

    def numbered(self, items: list[str], bold_leads: bool = True) -> None:
        """Numbered list; with bold_leads, text before the first ':' is bolded."""
        for index, item in enumerate(items, start=1):
            p = self._list_paragraph(f"{index}.\t")
            if bold_leads and ":" in item:
                head, _, tail = item.partition(":")
                r1 = p.add_run(f"{head}:")
                r1.font.name = FONT
                r1.font.size = Pt(10.5)
                r1.font.bold = True
                r2 = p.add_run(tail)
                r2.font.name = FONT
                r2.font.size = Pt(10.5)
            else:
                run = p.add_run(item)
                run.font.name = FONT
                run.font.size = Pt(10.5)

    def table(self, headers: list[str], rows: list[list[str]],
              widths: list[float] | None = None,
              header_fill: str = DREAM_HEX) -> None:
        t = self.doc.add_table(rows=1 + len(rows), cols=len(headers))
        _set_borders(t)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for j, header in enumerate(headers):
            cell = t.rows[0].cells[j]
            _shade(cell, header_fill)
            _cell_text(cell, header, bold=True, white=True)
        for i, row in enumerate(rows, start=1):
            for j, value in enumerate(row):
                _cell_text(t.rows[i].cells[j], value)
        if widths:
            for j, width in enumerate(widths):
                for row in t.rows:
                    row.cells[j].width = Cm(width)

    def kv_table(self, rows: list[tuple[str, str]], key_width: float = 4.5,
                 val_width: float = 12.0) -> None:
        """Two-column key-value table (the 'At a glance' block)."""
        t = self.doc.add_table(rows=len(rows), cols=2)
        _set_borders(t)
        for i, (key, value) in enumerate(rows):
            _shade(t.rows[i].cells[0], DREAM_HEX)
            _cell_text(t.rows[i].cells[0], key, bold=True, white=True)
            _cell_text(t.rows[i].cells[1], value)
            t.rows[i].cells[0].width = Cm(key_width)
            t.rows[i].cells[1].width = Cm(val_width)

    def save(self, out_path: Path | None = None) -> Path:
        out = out_path or Path(__file__).resolve().parent.parent / OUT_NAME
        target = str(out)
        # Windows MAX_PATH: docx saves fail past ~260 chars unless the path
        # carries the \\?\ extended-length prefix. Applied only when needed so
        # normal project paths stay untouched.
        if os.name == "nt" and len(target) > 255 and not target.startswith("\\\\?\\"):
            target = "\\\\?\\" + target
        self.doc.save(target)
        return out


def build_tor(header: str = "slim") -> Path:
    """REPLACE THIS BODY with the real ToR, section by section.

    Canonical order (see references/open-market-checklist.md):
    internal prep note, at-a-glance, how-to-apply, assignment in brief,
    sections 1-13, annexes.

    Defaults to slim because a ToR runs to several pages and the contact block
    is worth reading once, not on every page.
    """
    b = TorBuilder(header=header)
    b.title("Terms of Reference — [assignment title]", "[commissioning unit] — DRAFT v0.1")
    b.h1("Internal preparation note — REMOVE BEFORE PUBLICATION")
    b.bullets(["[unresolved item 1]", "[unresolved item 2]"])
    b.h1("At a glance")
    b.kv_table([
        ("Assignment", "[title]"),
        ("Commissioned by", "[entity]"),
        ("Maximum budget", "EUR [X], inclusive of all fees, costs and any applicable VAT"),
        ("Level of effort", "Proposed by the consultant: days per deliverable"),
        ("Duration", "[start] to [end]"),
    ])
    b.h1("How to apply")
    b.para("[deadline, submit-to address, what to submit, Q&A dates, interview dates]")
    return b.save()


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Build an IPPF EN-branded ToR docx.")
    parser.add_argument(
        "--header", choices=("slim", "full"), default="slim",
        help="slim = band plus logo (default, best for a multi-page ToR); "
             "full = keep the header contact block on every page",
    )
    args = parser.parse_args()
    print(build_tor(header=args.header))
