"""Checks for the office-review-pass CLI driver.

The driver is thin, so these checks are about the things a thin wrapper still
gets wrong: reading both edit-file shapes, refusing to write over the source,
and — the one that matters — reporting a search string that matched nothing as a
failure rather than as a quiet success. A review round that silently applied
none of its edits looks exactly like one that applied all of them.

Office behaviour itself is tested in the work folder's tests/test_officeops.py.

Run:  python test_review_pass.py
"""

from __future__ import annotations

import json
import subprocess
import sys
import tempfile
from pathlib import Path

DRIVER = Path(__file__).resolve().parent / "review_pass.py"

sys.path.insert(0, str(DRIVER.parent))
import review_pass  # noqa: E402

from ane_package.officeops import Checks, read_comments  # noqa: E402
from ane_package.officeops import CommentRequest, add_comments  # noqa: E402


def _run(*args: str) -> subprocess.CompletedProcess:
    return subprocess.run([sys.executable, str(DRIVER), *args],
                          capture_output=True, text=True)


def _fixture(tmp: Path) -> Path:
    from docx import Document

    document = Document()
    document.add_heading("Terms of reference", level=1)
    document.add_paragraph(
        "The consultant will submit an inception report within two weeks of signature."
    )
    document.add_paragraph(
        "Payment follows written acceptance of each deliverable by the unit."
    )
    path = tmp / "driver-source.docx"
    document.save(str(path))
    return path


def test_edit_loading(checks: Checks, tmp: Path) -> None:
    mapping = tmp / "edits-dict.json"
    mapping.write_text(json.dumps({"two weeks": "ten working days"}), encoding="utf-8")
    checks.check(review_pass._load_edits(mapping) == [("two weeks", "ten working days")],
                 "edits: object form loads")

    listed = tmp / "edits-list.json"
    listed.write_text(json.dumps([{"old": "a", "new": "b"}, {"old": "c", "new": "d"}]),
                      encoding="utf-8")
    checks.check(review_pass._load_edits(listed) == [("a", "b"), ("c", "d")],
                 "edits: list form loads and keeps its order")

    empty_new = tmp / "edits-delete.json"
    empty_new.write_text(json.dumps([{"old": "strike this"}]), encoding="utf-8")
    checks.check(review_pass._load_edits(empty_new) == [("strike this", "")],
                 "edits: a missing 'new' reads as a deletion")

    broken = tmp / "edits-broken.json"
    broken.write_text(json.dumps([{"new": "b"}]), encoding="utf-8")
    try:
        review_pass._load_edits(broken)
        checks.check(False, "edits: an entry with no 'old' key is rejected")
    except ValueError:
        checks.check(True, "edits: an entry with no 'old' key is rejected")


def test_read_command(checks: Checks, tmp: Path) -> None:
    source = _fixture(tmp)
    commented = add_comments(
        source,
        [CommentRequest(match="inception report within two weeks",
                        text="Name the acceptance criterion.", author="Ane Gasser")],
        out_path=tmp / "driver-commented.docx",
    )
    checks.check(len(read_comments(commented)) == 1, "read: fixture carries one comment")

    out = tmp / "review.md"
    result = _run("read", str(commented), "--out", str(out))
    checks.check(result.returncode == 0, f"read: exits 0 (stderr={result.stderr[:120]})")
    checks.check(out.is_file(), "read: rendering written to the named file")
    if out.is_file():
        text = out.read_text(encoding="utf-8")
        checks.check("Name the acceptance criterion." in text,
                     "read: the comment travels into the rendering")
        checks.check("Terms of reference" in text, "read: body text present")
    checks.check("3 blocks, 1 comment shown" in result.stdout,
                 f"read: block and comment counts reported to the caller, singular "
                 f"not '1 comments' (got {result.stdout.strip()[-70:]!r})")

    to_stdout = _run("read", str(commented))
    checks.check("Terms of reference" in to_stdout.stdout,
                 "read: renders to stdout when no --out is given")


def test_track_command(checks: Checks, tmp: Path) -> None:
    source = _fixture(tmp)
    edits = tmp / "track-edits.json"
    edits.write_text(json.dumps({"within two weeks": "within ten working days"}),
                     encoding="utf-8")
    out = tmp / "driver-tracked.docx"

    result = _run("track", str(source), "--edits", str(edits),
                  "--author", "Ane Gasser", "--out", str(out))
    checks.check(result.returncode == 0, f"track: exits 0 (stderr={result.stderr[:120]})")
    checks.check(out.is_file(), "track: marked-up copy written")
    checks.check(read_comments(source) == [] and source.is_file(),
                 "track: the source document is left alone")

    listing = _run("revisions", str(out))
    checks.check("within ten working days" in listing.stdout,
                 "revisions: the insertion is listed")
    checks.check("Ane Gasser" in listing.stdout, "revisions: author attribution shown")
    checks.check("2 revisions" in listing.stdout,
                 f"revisions: one insertion and one deletion "
                 f"(got {listing.stdout.strip().splitlines()[-1]!r})")

    # The check this driver exists for: a search string that matched nothing has
    # to fail loudly, because the written file looks identical either way.
    missing = tmp / "track-missing.json"
    missing.write_text(json.dumps({"a phrase that is not in the document": "x"}),
                       encoding="utf-8")
    result = _run("track", str(source), "--edits", str(missing),
                  "--out", str(tmp / "driver-missed.docx"))
    checks.check(result.returncode == 1,
                 f"track: a search string matching nothing exits non-zero "
                 f"(got {result.returncode})")
    checks.check("matched NOTHING" in result.stderr,
                 "track: the unmatched string is named on stderr")

    refused = _run("track", str(source), "--edits", str(edits), "--out", str(source))
    checks.check(refused.returncode == 2,
                 "track: refuses to write the marked-up copy over the source")


def test_verify_command(checks: Checks, tmp: Path) -> None:
    source = _fixture(tmp)
    result = _run("verify", str(source))
    checks.check("word count computed from the written file" in result.stdout,
                 "verify: word count asserted on the file")
    checks.check("stranded" in result.stdout, "verify: stranded hyperlinks asserted")

    # A plain python-docx document has no logo, so the branded assertion must fail
    # rather than pass by default — that silent pass is how unbranded ToRs shipped.
    branded = _run("verify", str(source), "--expect-branded")
    checks.check(branded.returncode != 0,
                 "verify: --expect-branded fails on an unbranded document")


def main() -> int:
    checks = Checks(title="office-review-pass driver")
    tmp = Path(tempfile.mkdtemp(prefix="review-pass-tests-"))

    test_edit_loading(checks, tmp)
    test_read_command(checks, tmp)
    test_track_command(checks, tmp)
    test_verify_command(checks, tmp)

    code = checks.report()
    print(f"artefacts: {tmp}")
    return code


if __name__ == "__main__":
    sys.exit(main())
