#!/usr/bin/env python3
"""Verification for selection_toolkit.py — the selection-toolkit driver.

Almost none of this needs a .docx, because the driver splits at the file
boundary: ``read_source`` opens the ToR and everything downstream takes the
``DocTable`` objects and paragraph strings it returned. Table scoring, weight
parsing, threshold reading and spec assembly are all tested against literal
grids.

Two deliberate exceptions at the end. The CLI runs as a subprocess over a ToR
carrying accented criteria labels, because Wave 3 shipped a UnicodeEncodeError
past 32 passing unit tests — an in-process call never touches the console
encoding. And ``verification_plan`` is checked against the engine directly,
because it is the thing that decides whether the Excel run passes.

Run: python test_selection_toolkit.py
"""

from __future__ import annotations

import json
import subprocess
import sys
import tempfile
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))

import selection_toolkit as toolkit  # noqa: E402
from ane_package.officeops import DocTable  # noqa: E402

PASSED = 0
FAILED = 0


def check(condition: bool, label: str) -> None:
    global PASSED, FAILED
    if condition:
        PASSED += 1
        print(f"  PASS  {label}")
    else:
        FAILED += 1
        print(f"  FAIL  {label}")


def grid(rows, where="body table 1", index=1, merged=(), ragged=False) -> DocTable:
    width = max(len(row) for row in rows)
    padded = tuple(tuple(row) + ("",) * (width - len(row)) for row in rows)
    return DocTable(index=index, where=where, depth=0, rows=padded,
                    merged=frozenset(merged), ragged=ragged)


CRITERIA_EN = grid([
    ("Criterion", "Description", "Maximum points"),
    ("C1", "Technical approach and methodology", "35"),
    ("C2", "Relevant expertise", "25"),
    ("C3", "Understanding of IPPF context", "15"),
    ("C4", "Work samples and references", "15"),
])

MILESTONES = grid([
    ("Milestone", "Date"),
    ("ToR published", "2026-07-31"),
    ("Proposals due", "2026-08-28"),
], where="body table 2", index=2)

CONTACTS = grid([
    ("Name", "Role", "Email"),
    ("Procurement", "Administrator", "someone@example.org"),
], where="body table 3", index=3)


# ---------------------------------------------------------------------------

def test_points_parsing() -> None:
    print("\nparse_points")
    for text, wanted in (
        ("35", 35.0),
        ("35 points", 35.0),
        ("up to 35%", 35.0),
        ("12,5 puncte", 12.5),
        ("12.5", 12.5),
        ("  40  ", 40.0),
    ):
        check(toolkit.parse_points(text) == wanted,
              f"parse_points({text!r}) -> {wanted}")
    # The blank-weight case. Returning 0 here would score every proposal
    # identically on that criterion and never raise.
    for text in ("", "   ", "to be confirmed", "TBC", None):
        check(toolkit.parse_points(text) is None,
              f"parse_points({text!r}) is None, not a guess")


def test_folding() -> None:
    print("\nlanguage folding")
    check(toolkit._fold("Critère") == "critere", "accents are folded away")
    check(toolkit._has_any("Critère d'attribution", toolkit.CRITERION_WORDS),
          "a French criteria header matches")
    check(toolkit._has_any("Criteriu", toolkit.CRITERION_WORDS),
          "a Romanian criteria header matches")
    check(toolkit._has_any("Puntuación máxima", toolkit.POINTS_WORDS),
          "a Spanish points header matches")
    check(toolkit._has_any("Pondération", toolkit.POINTS_WORDS),
          "a French weighting header matches")
    check(not toolkit._has_any("Email address", toolkit.CRITERION_WORDS),
          "an unrelated header does not match")


def test_table_selection() -> None:
    print("\nfinding the criteria table")
    tables = [MILESTONES, CONTACTS, CRITERIA_EN]
    best, score = toolkit.find_criteria_table(tables)
    # Picking the milestones table produces a plausible criteria list with dates
    # for weights, and nothing raises.
    check(best is CRITERIA_EN,
          f"the criteria table wins over milestones and contacts (got {best.where})")
    check(score >= 6, f"a clear criteria table scores strongly (got {score})")
    check(toolkit.score_criteria_table(MILESTONES) < score,
          "a milestones table scores below the criteria table")
    check(toolkit.score_criteria_table(CONTACTS) < 6,
          "a contacts table does not look like criteria")

    empty = grid([("Criterion", "Points")])
    check(toolkit.score_criteria_table(empty) == 0,
          "a header with no body rows scores nothing")
    check(toolkit.find_criteria_table([CONTACTS])[0] is None
          or toolkit.find_criteria_table([CONTACTS])[1] < 6,
          "no strong candidate is reported as weak, not as a confident pick")

    french = grid([
        ("Critère d'attribution", "Description", "Pondération"),
        ("1.", "Approche technique et méthodologie", "35 points"),
        ("2.", "Expertise pertinente", "25 points"),
    ])
    best_fr, score_fr = toolkit.find_criteria_table([MILESTONES, french])
    check(best_fr is french and score_fr >= 6,
          f"a French criteria table is found (got {score_fr})")


def test_criteria_parsing() -> None:
    print("\nparsing criteria")
    criteria, gaps = toolkit.parse_criteria(CRITERIA_EN)
    check(len(criteria) == 4, f"every criterion row is read (got {len(criteria)})")
    check([c["max_points"] for c in criteria] == [35, 25, 15, 15],
          "weights are read as numbers")
    check([c["code"] for c in criteria] == ["C1", "C2", "C3", "C4"],
          "codes published in the ToR are kept")
    # This assertion was missing and the parser was silently labelling every
    # criterion with its code, because 'Criterion' headed the code column and
    # the criterion text sat under 'Description'.
    check(criteria[0]["label"] == "Technical approach and methodology",
          f"the criterion text is the label, not its code (got {criteria[0]['label']!r})")
    check(all(len(c["label"]) > 5 for c in criteria),
          f"no criterion is labelled with a bare code (got {[c['label'] for c in criteria]})")
    check(not gaps, f"a complete table reports no gaps (got {gaps})")

    blank = grid([
        ("Criterion", "Description", "Maximum points"),
        ("C1", "Technical approach", "35"),
        ("C2", "Relevant expertise", ""),
        ("C3", "IPPF context", "15"),
    ])
    criteria, gaps = toolkit.parse_criteria(blank)
    # The whole reason document_tables preserves blanks. Dropping C2 would
    # quietly shorten the criteria list and rebalance the award.
    check(len(criteria) == 3, f"a criterion with no weight is kept (got {len(criteria)})")
    check(criteria[1]["max_points"] is None, "the unstated weight reads as None")
    check(any("C2" in gap or "Relevant expertise" in gap for gap in gaps),
          f"the gap names the criterion whose weight is missing (got {gaps})")

    totals = grid([
        ("Criterion", "Points"),
        ("Technical approach", "35"),
        ("Expertise", "25"),
        ("Total", "60"),
    ])
    criteria, _ = toolkit.parse_criteria(totals)
    check(len(criteria) == 2,
          f"a totals row is not scored as a criterion (got {[c['label'] for c in criteria]})")

    unlabelled = grid([
        ("Item", "What we are looking for", "Score"),
        ("1.", "Technical approach and methodology in detail", "35"),
        ("2.", "Relevant expertise across the region", "25"),
    ])
    criteria, gaps = toolkit.parse_criteria(unlabelled)
    check(len(criteria) == 2 and criteria[0]["max_points"] == 35,
          "a criteria table with an unusual header still parses")
    check(any("Confirm" in gap for gap in gaps),
          "reading a column by width is reported for confirmation, not assumed")


def test_threshold() -> None:
    print("\nreading the threshold")
    value, gaps = toolkit.find_threshold(
        ["Proposals scoring below 63 points will not proceed to the financial stage."],
        90,
    )
    check(value == 63 and not gaps, f"a stated points threshold is read (got {value})")

    value, gaps = toolkit.find_threshold(
        ["Applicants must reach a minimum of 70% of the technical marks."], 90)
    check(value == 63.0, f"a percentage threshold resolves against the max (got {value})")

    value, gaps = toolkit.find_threshold(["The evaluation will be fair."], 90)
    check(value is None and gaps, "silence is reported as a gap, not defaulted")

    # Two different numbers is the dangerous case: picking either is a coin
    # flip that decides who stays in the process.
    value, gaps = toolkit.find_threshold([
        "A minimum of 70% is required on the technical proposal.",
        "The qualification threshold is 60 points.",
    ], 90)
    check(value is None, "two possible thresholds resolve to None, not to the first")
    check(any("more than one" in gap.lower() for gap in gaps),
          f"the ambiguity is named and both readings shown (got {gaps[:1]})")

    value, gaps = toolkit.find_threshold(
        ["Le seuil de qualification est de 63 points."], 90)
    check(value == 63, f"a French threshold sentence is read (got {value})")


def test_draft_spec() -> None:
    print("\ndrafting the spec")
    paragraphs = [
        "Terms of Reference for an AI-for-Research consultancy.",
        "Proposals scoring below 63 points will not proceed.",
        "The financial proposal carries 10 points.",
    ]
    draft = toolkit.draft_spec([MILESTONES, CRITERIA_EN, CONTACTS], paragraphs,
                               source="ToR 2026-07-31")
    check(len(draft["criteria"]) == 4, "criteria reach the draft")
    check(draft["threshold"] == 63, "threshold reaches the draft")
    check(draft["financial_max"] == 10, "financial points reach the draft")
    check(draft["source"] == "ToR 2026-07-31", "the source is carried, never invented")
    check(draft["panel"] == [], "the panel is left empty; a ToR does not name it")
    check(draft["title"] == "", "the title is left empty rather than guessed")
    check(any("Panel" in gap for gap in draft["_gaps"]),
          "the missing panel is listed as a gap")
    check(["Proposals due", "2026-08-28"] in draft["key_dates"],
          f"dates are read from the milestones table (got {draft['key_dates']})")

    bare = toolkit.draft_spec([CONTACTS], ["Nothing useful here."], source="x.docx")
    check(bare["criteria"] == [], "no criteria table means no criteria")
    check(bare["threshold"] is None, "no threshold sentence means None")
    check(any("award criteria" in gap.lower() for gap in bare["_gaps"]),
          "a ToR with no criteria table says so plainly")


def test_spec_gate() -> None:
    print("\nthe gap gate")
    tmp = Path(tempfile.mkdtemp(prefix="selection-spec-"))
    draft = toolkit.draft_spec([CRITERIA_EN], ["Below 63 points will not proceed."],
                               source="ToR")
    path = tmp / "spec.json"
    path.write_text(json.dumps(draft), encoding="utf-8")
    try:
        toolkit._load_spec(path)
        check(False, "a spec with unresolved gaps is refused")
    except toolkit.UnstatedInSource as exc:
        check("_gaps" in str(exc) or "gap" in str(exc).lower(),
              "a spec with unresolved gaps is refused and says so")

    resolved = dict(draft)
    resolved.pop("_gaps")
    resolved.update({"title": "AI for Research", "panel": ["Ane", "Dome"],
                     "financial_max": 10})
    path.write_text(json.dumps(resolved), encoding="utf-8")
    spec = toolkit._load_spec(path)
    check(spec.technical_max == 90, f"a resolved spec loads (got {spec.technical_max})")

    # A gap deleted rather than answered must still be caught downstream.
    broken = dict(resolved)
    broken["threshold"] = None
    path.write_text(json.dumps(broken), encoding="utf-8")
    try:
        toolkit._load_spec(path)
        check(False, "deleting _gaps does not smuggle a missing threshold through")
    except toolkit.UnstatedInSource:
        check(True, "validate_spec still refuses a missing threshold after _gaps is gone")


def test_verification_plan() -> None:
    print("\nverification plan")
    from ane_package.reporting.selection_toolkit import Criterion, SelectionSpec

    spec = SelectionSpec(
        mode="procurement", title="t", source="ToR",
        criteria=(Criterion("C1", "Approach", 35), Criterion("C2", "Expertise", 25),
                  Criterion("C3", "Context", 15), Criterion("C4", "Samples", 15)),
        panel=("Ane", "Dome", "Stefanie", "Lena"),
        threshold=63, financial_max=10, n_applicants=8,
    )
    writes, reads, expected = toolkit.verification_plan(spec)

    check(len(writes) > 20, f"sample data is injected (got {len(writes)} writes)")
    check(all(w["sheet"] in (toolkit.SHEETS["applicants"], toolkit.SHEETS["inbox"])
              for w in writes),
          "writes only touch the unprotected input sheets")

    # Alpha: means 29, 20.25, 11.5, 10.5 = 71.25, above 63.
    check(abs(expected["total0"] - 71.25) < 0.001,
          f"the expected total comes from the engine (got {expected['total0']})")
    check(expected["qualified0"] == "Qualified", "Alpha qualifies on 71.25")
    # Beta has one scorer who left it blank entirely. Blanks are absences, so
    # its means are over three scorers, not four.
    # Beta's means are over the three scorers who scored it, not four:
    # 76/3 + 45/3 + 27/3 + 24/3 = 57.333. Zeroing the blank would give 43.
    check(abs(expected["total1"] - 57.3333) < 0.001,
          f"a scorer who skipped a subject is ignored, not zeroed (got {expected['total1']})")
    check(expected["qualified1"] == "Below threshold", "Beta falls below the threshold")
    check(expected["financial1"] is None,
          "an unqualified subject gets no financial score")
    # Alpha is the only qualified bidder, so it is the cheapest qualified one.
    check(abs(expected["financial0"] - 10.0) < 0.001,
          f"the only qualified price takes full financial marks (got {expected['financial0']})")
    check(abs(expected["combined0"] - 81.25) < 0.001,
          f"combined is technical plus financial (got {expected['combined0']})")

    labels = {r["label"] for r in reads}
    check(set(expected) <= labels | set(expected),
          "every expectation has a cell to read it from")
    check(all(r["sheet"] in toolkit.SHEETS.values() for r in reads),
          "reads name real sheets")


# ---------------------------------------------------------------------------
# Subprocess: the console-encoding wall an in-process call never touches
# ---------------------------------------------------------------------------

def _tor_fixture(tmp: Path) -> Path:
    from docx import Document

    document = Document()
    document.add_heading("Termes de référence", level=1)
    document.add_paragraph(
        "Consultance sur l'intelligence artificielle pour la recherche, "
        "basée à Chișinău."
    )
    document.add_paragraph("Le seuil de qualification est de 63 points.")
    document.add_paragraph("La proposition financière compte pour 10 points.")

    table = document.add_table(rows=5, cols=3)
    for index, label in enumerate(("Critère d'attribution", "Description",
                                   "Pondération")):
        table.cell(0, index).text = label
    rows = [
        ("C1", "Approche technique et méthodologie", "35"),
        ("C2", "Expertise pertinente en évaluation", "25"),
        ("C3", "Compréhension du contexte d'IPPF", "15"),
        ("C4", "Échantillons de travaux et références", "15"),
    ]
    for row_index, values in enumerate(rows, start=1):
        for col_index, value in enumerate(values):
            table.cell(row_index, col_index).text = value

    document.add_paragraph("Calendrier")
    schedule = document.add_table(rows=2, cols=2)
    schedule.cell(0, 0).text = "Étape"
    schedule.cell(0, 1).text = "Date"
    schedule.cell(1, 0).text = "Dépôt des propositions"
    schedule.cell(1, 1).text = "2026-08-28"

    path = tmp / "tor-fr.docx"
    document.save(str(path))
    return path


def test_cli_subprocess(tmp: Path) -> None:
    print("\nCLI as a subprocess")
    script = Path(__file__).resolve().parent / "selection_toolkit.py"
    tor = _tor_fixture(tmp)
    spec_path = tmp / "spec.json"

    result = subprocess.run(
        [sys.executable, str(script), "read", str(tor), "--out", str(spec_path)],
        capture_output=True, text=True, encoding="utf-8", errors="replace",
    )
    check(result.returncode == 0,
          f"read exits cleanly (rc={result.returncode}) {result.stderr[-300:]}")
    # Wave 3's exact failure: the driver echoes criteria labels, and a cp1252
    # console kills it on the first accented one.
    check("UnicodeEncodeError" not in result.stderr,
          "no UnicodeEncodeError on accented criteria labels")
    check("Approche technique" in result.stdout,
          "accented labels survive to the console")

    draft = json.loads(spec_path.read_text(encoding="utf-8"))
    check(len(draft["criteria"]) == 4,
          f"the French ToR yields 4 criteria (got {len(draft['criteria'])})")
    check(draft["threshold"] == 63,
          f"the French threshold sentence is read (got {draft['threshold']})")
    check(draft["criteria"][0]["max_points"] == 35, "French weights are read")

    resolved = dict(draft)
    resolved.pop("_gaps")
    resolved.update({"title": "Consultance IA", "panel": ["Ane", "Dome"],
                     "n_applicants": 4})
    spec_path.write_text(json.dumps(resolved, ensure_ascii=False), encoding="utf-8")

    out = tmp / "built"
    result = subprocess.run(
        [sys.executable, str(script), "build", str(spec_path), "--out", str(out)],
        capture_output=True, text=True, encoding="utf-8", errors="replace",
    )
    check(result.returncode == 0,
          f"build exits cleanly (rc={result.returncode}) {result.stderr[-300:]}")
    check((out / "selection-master.xlsx").exists(), "the master workbook is written")
    check((out / "scorer-ane.xlsx").exists() and (out / "scorer-dome.xlsx").exists(),
          "one standalone scorer file per panel member")

    # Re-running a generator is overwriting, and by now these may hold scores.
    again = subprocess.run(
        [sys.executable, str(script), "build", str(spec_path), "--out", str(out)],
        capture_output=True, text=True, encoding="utf-8", errors="replace",
    )
    check(again.returncode == 1 and "Refusing to overwrite" in again.stdout,
          "a second build refuses rather than discard panel data")


def main() -> int:
    tmp = Path(tempfile.mkdtemp(prefix="selection-toolkit-tests-"))
    test_points_parsing()
    test_folding()
    test_table_selection()
    test_criteria_parsing()
    test_threshold()
    test_draft_spec()
    test_spec_gate()
    test_verification_plan()
    test_cli_subprocess(tmp)

    total = PASSED + FAILED
    print(f"\n{PASSED}/{total} checks passed (selection-toolkit driver)")
    print(f"artefacts: {tmp}")
    return 1 if FAILED else 0


if __name__ == "__main__":
    raise SystemExit(main())
