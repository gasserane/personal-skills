"""Add anchored margin comments to a copy of a supplier offer (never the original).

Usage:
    python add_offer_comments.py <offer.docx> <comments.json>

comments.json is a list of objects:
    [
      {
        "match": "5+ consecutive words uniquely identifying the target paragraph",
        "text": "The comment body.",
        "author": "Ane Gasser (MEL review)"   # or "Ane Gasser PERSONAL"
      }
    ]

Writes <offer>_COMMENTS.docx next to the original. Errors (rather than guesses)
when a match string hits zero or multiple paragraphs.
"""
import json
import shutil
import sys
from pathlib import Path

import docx


def main(offer_path: str, comments_path: str) -> None:
    offer = Path(offer_path)
    comments = json.loads(Path(comments_path).read_text(encoding="utf-8"))

    target = offer.with_name(offer.stem + "_COMMENTS.docx")
    shutil.copy2(offer, target)
    doc = docx.Document(str(target))

    paragraphs = [p for p in doc.paragraphs if p.text.strip()]
    errors = []
    for i, c in enumerate(comments):
        matches = [p for p in paragraphs if c["match"] in p.text]
        if len(matches) != 1:
            errors.append(
                f"comment {i}: match string hit {len(matches)} paragraphs "
                f"(need exactly 1): {c['match'][:60]!r}"
            )
            continue
        para = matches[0]
        author = c.get("author", "Ane Gasser (MEL review)")
        initials = "".join(w[0] for w in author.split()[:2]).upper()
        doc.add_comment(para.runs or [para.add_run("")], c["text"],
                        author=author, initials=initials)

    if errors:
        for e in errors:
            print(f"ERROR: {e}", file=sys.stderr)
        sys.exit(1)

    doc.save(str(target))
    print(f"OK: {len(comments)} comment(s) written to {target}")


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print(__doc__, file=sys.stderr)
        sys.exit(2)
    main(sys.argv[1], sys.argv[2])
