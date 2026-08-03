"""Checks on the driver itself, not on Office.

Office behaviour is covered by ``tests/test_officeops.py`` in the work folder
(123 checks). What is tested here is the workflow logic this skill owns: the
change list, the artefact heuristics, and the generator guard.

The guard gets a real test rather than an inspection — a temporary generator is
written, archived, and then actually run as a subprocess. A refusal that only
exists in a docstring is the failure this whole mode exists to prevent, so
asserting on the source text would be testing the wrong thing.

    python test_office_repair.py
"""

from __future__ import annotations

import subprocess
import sys
import tempfile
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))

from office_repair import (  # noqa: E402
    ARCHIVE_MARKER,
    archive_generator,
    diff_blocks,
    suspect_artefacts,
)
from ane_package.officeops import Checks  # noqa: E402


def test_diff(checks: Checks) -> None:
    before = ["Title", "First paragraph.", "Second paragraph.", "Closing."]
    after = ["Title", "First paragraph, revised.", "Second paragraph.",
             "A new sentence.", "Closing."]
    changes = diff_blocks(before, after)

    ops = [change["op"] for change in changes]
    checks.check("replace" in ops, f"diff: a rewritten block is a replace (got {ops})")
    checks.check("insert" in ops, f"diff: an added block is an insert (got {ops})")
    checks.check(all(change["op"] != "equal" for change in changes),
                 "diff: unchanged blocks are not reported")
    checks.check(len(changes) == 2,
                 f"diff: two edits produce two changes, not five (got {len(changes)})")

    checks.check(diff_blocks(before, before) == [],
                 "diff: an unchanged file produces an empty change list")

    # A block appearing earlier shifts every index after it. Pairing on position
    # would call the whole document changed.
    shifted = ["New opening."] + before
    checks.check(len(diff_blocks(before, shifted)) == 1,
                 "diff: an insertion at the top does not report the rest as changed")


def test_suspect_artefacts(checks: Checks) -> None:
    before = ["The workshop ran in Chișinău.", "Contact: https://ippfen.org/contact",
              "We agreed the “house style” applies.", "A genuinely rewritten line."]
    after = ["The workshop ran in Chisinau.", "Contact: https://ippfen.org/ contact",
             'We agreed the "house style" applies.',
             "This line says something completely different now."]
    flags = suspect_artefacts(diff_blocks(before, after), before, after)
    kinds = {flag["kind"] for flag in flags}

    checks.check("accents dropped" in kinds,
                 f"suspect: a diacritic strip is flagged (got {kinds})")
    checks.check("punctuation autocorrected" in kinds,
                 f"suspect: a smart-quote flip is flagged (got {kinds})")
    checks.check("space inside a hyperlink" in kinds,
                 f"suspect: a space inside a URL is flagged (got {kinds})")
    checks.check(all("completely different" not in str(flag["after"]) for flag in flags),
                 "suspect: a real rewrite is left alone")
    checks.check(all(flag["why"] for flag in flags),
                 "suspect: every flag says why it was raised")

    clean_before = ["One.", "Two."]
    clean_after = ["One.", "Two.", "Three, added deliberately."]
    checks.check(suspect_artefacts(diff_blocks(clean_before, clean_after),
                                   clean_before, clean_after) == [],
                 "suspect: a deliberate addition raises nothing")

    # The 2026-07-22 red flag: a redesign that leaves titles and drops the body.
    full = [f"Body block {index}" for index in range(20)]
    titles = ["Body block 0", "Body block 1"]
    collapse = suspect_artefacts(diff_blocks(full, titles), full, titles)
    checks.check(any(flag["kind"] == "most of the text is gone" for flag in collapse),
                 "suspect: a title-only layer after a redesign is caught structurally")


def _generator(tmp: Path, target: Path) -> Path:
    """A throwaway generator that overwrites its output, like every real one."""
    script = tmp / "gen_note.py"
    script.write_text(
        "#!/usr/bin/env python\n"
        "from pathlib import Path\n"
        f"Path(r'{target}').write_text('REGENERATED', encoding='utf-8')\n"
        "print('wrote the note')\n",
        encoding="utf-8",
    )
    return script


def test_archive_generator(checks: Checks, tmp: Path) -> None:
    canonical = tmp / "note.txt"
    canonical.write_text("ANE'S HAND-EDITED VERSION", encoding="utf-8")
    script = _generator(tmp, canonical)

    # Before archiving, the generator does exactly what it is supposed to do.
    subprocess.run([sys.executable, str(script)], capture_output=True, check=False)
    checks.check(canonical.read_text(encoding="utf-8") == "REGENERATED",
                 "archive: an unguarded generator overwrites the canonical file")

    canonical.write_text("ANE'S HAND-EDITED VERSION", encoding="utf-8")
    checks.check(archive_generator(script, canonical) == "archived",
                 "archive: the guard is written")
    source = script.read_text(encoding="utf-8")
    checks.check(ARCHIVE_MARKER in source, "archive: the file carries the marker")
    checks.check(source.startswith("#!/usr/bin/env python"),
                 "archive: a shebang stays on line 1")
    checks.check(archive_generator(script, canonical) == "already archived",
                 "archive: running it twice changes nothing")

    # The claim under test: running it refuses, and the canonical file survives.
    run = subprocess.run([sys.executable, str(script)], capture_output=True,
                         text=True, check=False)
    checks.check(run.returncode == 2,
                 f"archive: the archived generator exits non-zero (got {run.returncode})")
    checks.check("ARCHIVED" in run.stderr,
                 "archive: the refusal explains itself on stderr")
    checks.check(canonical.read_text(encoding="utf-8") == "ANE'S HAND-EDITED VERSION",
                 "archive: a plain run leaves the canonical file untouched")

    # --force lets the body run, and the guard puts the canonical file back.
    forced = subprocess.run([sys.executable, str(script), "--force"],
                            capture_output=True, text=True, check=False)
    checks.check("wrote the note" in forced.stdout,
                 "archive: --force lets the generator body execute")
    checks.check(canonical.read_text(encoding="utf-8") == "ANE'S HAND-EDITED VERSION",
                 "archive: --force still restores the canonical file byte-for-byte")
    checks.check("RESTORED" in forced.stderr,
                 "archive: the restore is reported rather than done silently")


def test_cli(checks: Checks, tmp: Path) -> None:
    driver = Path(__file__).resolve().parent / "office_repair.py"

    run = subprocess.run([sys.executable, str(driver), "--help"],
                         capture_output=True, text=True, check=False)
    checks.check(run.returncode == 0, "cli: --help exits zero")
    for command in ("scan", "repair", "verify", "diff", "archive"):
        checks.check(command in run.stdout, f"cli: {command} is a documented command")

    edits = tmp / "edits.json"
    edits.write_text('{"not": "a list"}', encoding="utf-8")
    book = tmp / "missing.xlsx"
    book.write_bytes(b"")
    bad = subprocess.run([sys.executable, str(driver), "repair", str(book),
                          "--edits", str(edits)], capture_output=True, text=True,
                         check=False)
    checks.check(bad.returncode == 2,
                 f"cli: an edits file of the wrong shape is refused (got {bad.returncode})")

    missing = subprocess.run([sys.executable, str(driver), "archive",
                              str(tmp / "nope.py"), "--canonical", str(tmp)],
                             capture_output=True, text=True, check=False)
    checks.check(missing.returncode == 2,
                 "cli: archiving a generator that does not exist is refused")

    # Only a real subprocess reproduces this: in-process calls never touch the
    # console encoding, so the unit tests above all passed while the shipped
    # command died on the first accented word it was asked to compare.
    from docx import Document

    for name, third in (("gen.docx", "The workshop ran in Chișinău."),
                        ("edited.docx", "The workshop ran in Chisinau.")):
        document = Document()
        document.add_paragraph("A baseline paragraph that did not change.")
        document.add_paragraph(third)
        document.save(str(tmp / name))

    diff = subprocess.run(
        [sys.executable, str(driver), "diff", str(tmp / "edited.docx"),
         "--against", str(tmp / "gen.docx")],
        capture_output=True, text=True, encoding="utf-8", check=False,
    )
    checks.check(diff.returncode == 0,
                 f"cli: diff survives non-ASCII document text (exit {diff.returncode})")
    checks.check("UnicodeEncodeError" not in diff.stderr,
                 "cli: diff does not die on the console encoding")
    checks.check("accents dropped" in diff.stdout,
                 "cli: diff flags the diacritic strip end to end")
    checks.check("Chișinău" in diff.stdout,
                 "cli: the accented original is printed, not mangled — it is the evidence")


def main() -> int:
    checks = Checks(title="office-repair driver")
    tmp = Path(tempfile.mkdtemp(prefix="office-repair-tests-"))

    test_diff(checks)
    test_suspect_artefacts(checks)
    test_archive_generator(checks, tmp)
    test_cli(checks, tmp)

    code = checks.report()
    print(f"artefacts: {tmp}")
    return code


if __name__ == "__main__":
    sys.exit(main())
