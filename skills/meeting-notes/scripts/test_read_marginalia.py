"""Checks on the driver, not on Office.

Reading comments out of a package is officeops' job and is covered by
``tests/test_officeops.py`` in the work folder. What is tested here is the
judgement this script owns: document order, the heading a comment sits under,
reply threading, and the date flag.

Most of it runs against hand-built ``ReviewBlock`` objects rather than a .docx,
which is only possible because ``assemble`` never touches the file. That is the
point of the split — logic that needs a document to test tends not to get
tested.

Two things do need a real run. A fixture .docx proves the officeops call and
this script's expectations still agree, and a SUBPROCESS proves the CLI can
print the accented text it exists to carry. Wave 3 shipped a
``UnicodeEncodeError`` past 32 passing unit tests because an in-process call
never touches the console encoding, so an in-process CLI check would repeat
that mistake rather than catch it.

    python test_read_marginalia.py
"""

from __future__ import annotations

import json
import subprocess
import sys
import tempfile
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))

from read_marginalia import (  # noqa: E402
    assemble,
    collect,
    date_mentions,
    render,
)
from ane_package.officeops import Checks, Comment, CommentRequest, ReviewBlock, add_comments  # noqa: E402


def _comment(cid: str, text: str, anchor: str = "", author: str = "Ane Gasser",
             date: str = "2026-07-24T10:00:00Z", parent: str | None = None,
             resolved: bool = False) -> Comment:
    return Comment(id=cid, author=author, initials="AG", date=date, text=text,
                   anchor=anchor, parent_id=parent, resolved=resolved)


def _block(index: int, text: str, style: str = "", comments=None,
           **kwargs) -> ReviewBlock:
    return ReviewBlock(index=index, text=text, style=style,
                       comments=list(comments or []), **kwargs)


def test_sections(checks: Checks) -> None:
    """A comment is reported under the heading it lives beneath, not the last one seen."""
    blocks = [
        _block(1, "Implementation plan", style="Title"),
        _block(2, "1. Budget", style="Heading1"),
        _block(3, "1.1 Follow-up funding", style="Heading2"),
        _block(4, "The follow-up funding bullet.",
               comments=[_comment("1", "Stef to discuss with Manuelle",
                                  anchor="The follow-up funding bullet.")]),
        _block(5, "2. Timeline", style="Heading1"),
        _block(6, "Workshop dates to confirm.",
               comments=[_comment("2", "Move this", anchor="Workshop dates to confirm.")]),
    ]
    data = assemble(blocks, Path("plan.docx"))

    first, second = data["comments"]
    checks.check(
        first["section"] == "Implementation plan > 1. Budget > 1.1 Follow-up funding",
        f"sections: the full heading path is carried (got {first['section']!r})",
    )
    checks.check(
        second["section"] == "Implementation plan > 2. Timeline",
        f"sections: a Heading1 pops the deeper Heading2 (got {second['section']!r})",
    )
    checks.check(first["block"] < second["block"],
                 "sections: comments come back in document order")
    checks.check(first["anchor"] == "The follow-up funding bullet.",
                 "sections: the anchor travels with the comment")


def test_anchor_is_never_lost(checks: Checks) -> None:
    """The 2026-07-24 failure: a comment that means nothing without its anchor."""
    blocks = [
        _block(1, "Follow-up funding: decide by September.",
               comments=[_comment("1", "Stef to discuss with Manuelle")]),
    ]
    data = assemble(blocks, Path("plan.docx"))
    checks.check(
        data["comments"][0]["anchor"] == "Follow-up funding: decide by September.",
        "anchor: a comment with no range falls back to its paragraph text, never empty",
    )
    checks.check("Follow-up funding" in render(data),
                 "anchor: the rendered note shows the anchor, not just the comment")


def test_replies(checks: Checks) -> None:
    blocks = [
        _block(1, "The clause.", comments=[
            _comment("1", "Change this to 12 October"),
            _comment("2", "Agreed", author="Stefanie", parent="1"),
        ]),
    ]
    data = assemble(blocks, Path("plan.docx"))
    checks.check(len(data["comments"]) == 1,
                 f"replies: a reply does not become its own thread (got {len(data['comments'])})")
    checks.check(data["thread_count"] == 1 and data["comment_count"] == 2,
                 f"replies: counts separate threads from comments "
                 f"({data['thread_count']} threads, {data['comment_count']} comments)")
    checks.check(data["comments"][0]["replies"][0]["author"] == "Stefanie",
                 "replies: the reply hangs off its parent")

    # A reply whose parent was filtered out still carries a decision. Dropping it
    # silently is worse than showing it unattached.
    orphan = assemble([_block(1, "The clause.", comments=[
        _comment("1", "Resolved point", resolved=True),
        _comment("2", "But note the date", author="Stefanie", parent="1"),
    ])], Path("plan.docx"), open_only=True)
    checks.check(len(orphan["comments"]) == 1,
                 f"replies: a reply outliving a filtered parent is kept, not dropped "
                 f"(got {len(orphan['comments'])})")


def test_resolved_filter(checks: Checks) -> None:
    blocks = [_block(1, "A clause.", comments=[
        _comment("1", "Open point"),
        _comment("2", "Done", resolved=True),
    ])]
    checks.check(len(assemble(blocks, Path("d.docx"))["comments"]) == 2,
                 "resolved: kept by default, because a resolved point is still a decision")
    checks.check(len(assemble(blocks, Path("d.docx"), open_only=True)["comments"]) == 1,
                 "resolved: --open-only drops them")


def test_date_flags(checks: Checks) -> None:
    """Dates get flagged, never corrected. The script cannot know which is right."""
    july = date_mentions("book it for 14-15 July", "2026-09-02T09:00:00Z")
    checks.check(len(july) == 1,
                 f"dates: '14-15 July' is one mention, not two numbers (got {len(july)})")
    checks.check(july[0]["phrase"] == "14-15 July",
                 f"dates: the whole range is captured (got {july[0]['phrase']!r})")
    checks.check(july[0]["backdated"],
                 "dates: a month earlier than the comment is called out — the 2026-07-24 shape")

    ahead = date_mentions("deliver by 12 October", "2026-07-24T09:00:00Z")
    checks.check(ahead and not ahead[0]["backdated"],
                 "dates: a forward date is flagged for confirmation but not as a slip")

    checks.check(date_mentions("increase this to 15 people", "2026-07-24T09:00:00Z") == [],
                 "dates: a bare number is not a date")
    checks.check(len(date_mentions("between 3 March and 4 April", "2026-07-24T09:00:00Z")) == 2,
                 "dates: two mentions in one comment are both reported")
    checks.check(any(m["phrase"] == "2026-10-12"
                     for m in date_mentions("due 2026-10-12", "2026-07-24T09:00:00Z")),
                 "dates: an ISO date is caught too")

    # Ane reviews in four languages. A guard that only reads English fails
    # silently on the documents where a slip is most likely.
    for phrase, comment_date, label in [
        ("reporter à 14 septembre", "2026-07-24T09:00:00Z", "French"),
        ("mover a 3 octubre", "2026-07-24T09:00:00Z", "Spanish"),
        ("mutăm pe 12 octombrie", "2026-07-24T09:00:00Z", "Romanian"),
    ]:
        checks.check(bool(date_mentions(phrase, comment_date)),
                     f"dates: a {label} month name is flagged too")
    checks.check(
        _month_name_used(date_mentions("14 mai", "2026-09-02T09:00:00Z")),
        "dates: the warning names the month in English however it was written",
    )
    checks.check(date_mentions("mai multe persoane", "2026-07-24T09:00:00Z") == [],
                 "dates: Romanian 'mai' meaning 'more' is not read as May")


def _month_name_used(mentions: list[dict]) -> bool:
    return bool(mentions) and "May" in mentions[0]["note"]

    rendered = render(assemble([_block(1, "A clause.", comments=[
        _comment("1", "move to 14-15 July", date="2026-09-02T09:00:00Z")])], Path("d.docx")))
    checks.check("Dates to confirm" in rendered,
                 "dates: the summary table appears when anything is flagged")
    checks.check("⚠️" in rendered, "dates: the flag is visible at the comment itself")


def test_tracked_anchor_warning(checks: Checks) -> None:
    """Anchor text on a tracked paragraph is the accepted state, not what was seen."""
    data = assemble([_block(1, "A revised clause.", has_insertions=True,
                            comments=[_comment("1", "Fine now")])], Path("d.docx"))
    checks.check(data["comments"][0]["tracked_anchor"],
                 "tracked: a comment on a changed paragraph is marked")
    checks.check("±tracked" in render(data) and "ACCEPTED" in render(data),
                 "tracked: the rendering warns that the anchor shows changes accepted")


def _fixture(tmp: Path) -> Path:
    """A real commented .docx, with accents, because the CLI has to print them."""
    from docx import Document

    document = Document()
    document.add_heading("Plan de mise en œuvre", level=1)
    document.add_paragraph("L'atelier se tiendra à Chișinău en septembre.")
    document.add_heading("Suivi", level=1)
    document.add_paragraph("Le financement de suivi reste à confirmer.")
    plain = tmp / "plan.docx"
    document.save(str(plain))

    return add_comments(plain, [
        CommentRequest(match="L'atelier se tiendra à Chișinău",
                       text="confirmer les dates, peut-être 14-15 July",
                       author="Ane Gasser"),
        CommentRequest(match="Le financement de suivi reste",
                       text="Stef to discuss with Manuelle", author="Ane Gasser"),
    ], out_path=tmp / "plan_reviewed.docx")


def test_real_document(checks: Checks, tmp: Path) -> None:
    reviewed = _fixture(tmp)
    data = collect(reviewed)

    checks.check(data["thread_count"] == 2,
                 f"document: both comments are read back (got {data['thread_count']})")
    sections = {item["section"] for item in data["comments"]}
    checks.check(sections == {"Plan de mise en œuvre", "Suivi"},
                 f"document: each comment lands under its own heading (got {sections})")
    checks.check(any("Chișinău" in item["anchor"] for item in data["comments"]),
                 "document: the anchor keeps its diacritics")
    checks.check(any(item["dates"] for item in data["comments"]),
                 "document: the date inside a comment is flagged end to end")


def test_cli(checks: Checks, tmp: Path) -> None:
    driver = Path(__file__).resolve().parent / "read_marginalia.py"
    reviewed = _fixture(tmp)

    # The regression that unit tests structurally cannot see. A Windows console
    # is cp1252; printing "Chișinău" through it raises unless stdout is
    # reconfigured, and only a real subprocess has a console to get wrong.
    run = subprocess.run([sys.executable, str(driver), str(reviewed)],
                         capture_output=True, check=False)
    stdout = run.stdout.decode("utf-8", errors="replace")
    checks.check(run.returncode == 0,
                 f"cli: a commented document exits 0 (got {run.returncode}: "
                 f"{run.stderr.decode('utf-8', errors='replace')[:200]})")
    checks.check("UnicodeEncodeError" not in run.stderr.decode("utf-8", errors="replace"),
                 "cli: printing accented document text does not blow up the console")
    checks.check("Chișinău" in stdout,
                 "cli: the diacritics survive to stdout rather than being mangled")
    checks.check("Dates to confirm" in stdout, "cli: the date flag reaches the output")

    as_json = subprocess.run([sys.executable, str(driver), str(reviewed), "--json"],
                             capture_output=True, check=False)
    parsed = json.loads(as_json.stdout.decode("utf-8"))
    checks.check(parsed["thread_count"] == 2, "cli: --json emits the same threads")

    out_file = tmp / "extract.md"
    subprocess.run([sys.executable, str(driver), str(reviewed), "--out", str(out_file)],
                   capture_output=True, check=False)
    checks.check(out_file.is_file() and "Chișinău" in out_file.read_text(encoding="utf-8"),
                 "cli: --out writes UTF-8 to disk")

    # A document with no comments is the commonest wrong input. It has to say so
    # rather than print an empty note that reads like a meeting with no decisions.
    from docx import Document
    bare = tmp / "bare.docx"
    Document().save(str(bare))
    empty = subprocess.run([sys.executable, str(driver), str(bare)],
                           capture_output=True, check=False)
    checks.check(empty.returncode == 1,
                 f"cli: an uncommented document fails loudly (got {empty.returncode})")
    checks.check("no comments" in empty.stderr.decode("utf-8", errors="replace"),
                 "cli: and says which document it looked at and what to try instead")

    missing = subprocess.run([sys.executable, str(driver), str(tmp / "nope.docx")],
                             capture_output=True, check=False)
    checks.check(missing.returncode == 2, "cli: a missing file is a usage error, not a crash")


def main() -> int:
    checks = Checks(title="meeting-notes marginalia driver")
    with tempfile.TemporaryDirectory() as raw:
        tmp = Path(raw)
        test_sections(checks)
        test_anchor_is_never_lost(checks)
        test_replies(checks)
        test_resolved_filter(checks)
        test_date_flags(checks)
        test_tracked_anchor_warning(checks)
        test_real_document(checks, tmp)
        test_cli(checks, tmp)
    return checks.report()


if __name__ == "__main__":
    raise SystemExit(main())
