#!/usr/bin/env python3
"""Verification for extract_deck.py — the localise deck-export mode.

Most of what follows needs no .docx, because the driver splits at the file
boundary: ``collect`` opens the document and everything downstream takes the
objects it returned. That is what makes slide numbering, chrome detection,
anchor collisions, comment merging and worklist rows cheap to cover.

The end of the file is the exception, and it is deliberate. Wave 3 shipped a
UnicodeEncodeError past 32 passing unit tests because an in-process call never
touches the console encoding, so the CLI is exercised as a subprocess over a
fixture carrying accented slide titles.

Run: python test_extract_deck.py
"""

from __future__ import annotations

import json
import subprocess
import sys
import tempfile
import zipfile
from dataclasses import dataclass
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))

import extract_deck as deck  # noqa: E402

PASSED = 0
FAILED = 0


def check(condition: bool, label: str) -> None:
    global PASSED, FAILED
    if condition:
        PASSED += 1
        print(f"  PASS  {label}")
    else:
        FAILED += 1
        print(f"  FAIL  {label}")


@dataclass
class FakeImage:
    """Stands in for officeops.SlideImage, which is all assemble() needs."""

    index: int
    anchor: str
    digest: str
    anchor_index: int = 0
    part: str = "media/image.bin"
    image_format: str = "jpeg"
    size: int = 40960
    path: Path | None = None


@dataclass
class FakeCheck:
    is_deck_export: bool = True
    reason: str = "test fixture"
    main_part: str = "word/document2.xml"


# --------------------------------------------------------------------------

def test_slide_numbering() -> None:
    print("\nslide numbering — every language localise supports")
    # The guard that matters: the deck under review is not in English.
    for text, expected in [
        ("Slide 7", 7),
        ("Diapositive 7", 7),          # French
        ("Diapositiva 7", 7),          # Spanish
        ("Diapozitiv 7", 7),           # Romanian
        ("Слайд 7", 7),                # Russian
        ("Diapositive 12 : Les indicateurs", 12),
        ("Slide-3", 3),
        ("Diapositive — 4", 4),
    ]:
        check(deck.slide_number(text) == expected,
              f"slide_number({text!r}) == {expected} (got {deck.slide_number(text)})")

    check(deck.slide_number("Introduction") is None,
          "a title with no slide label yields no number")
    check(deck.slide_number("") is None, "an empty anchor yields no number")
    # "mai" is Romanian for "more" and French for May. Neither is a slide label,
    # and neither should be read as one.
    check(deck.slide_number("mai multe informații") is None,
          "a word that merely starts like a label does not match")

    check(deck.is_label_only("Diapositive 4"), "a bare French slide label is label-only")
    check(deck.is_label_only("Calque 2"), "a French layer label is label-only")
    check(deck.is_label_only("Слой"), "a Russian layer label is label-only")
    check(not deck.is_label_only("Diapositive 4 : Les droits"),
          "a label carrying a title is not label-only")


def test_chrome_and_collisions() -> None:
    print("\nchrome and anchor collisions")
    images = [
        FakeImage(1, "Diapositive 1", "aaa"),
        FakeImage(2, "Diapositive 2", "logo"),
        FakeImage(3, "Diapositive 3", "bbb"),
        FakeImage(4, "Diapositive 4", "logo"),
        FakeImage(5, "Diapositive 5", "logo"),
    ]
    chrome = deck.find_chrome(images)
    check(chrome == {"logo"}, f"a digest on 3+ slides is chrome (got {chrome})")

    twice = [FakeImage(1, "a", "x"), FakeImage(2, "b", "x")]
    check(deck.find_chrome(twice) == set(),
          "two slides sharing an image is a coincidence, not chrome")

    repeated = [
        FakeImage(1, "Diapositive 1", "a"),
        FakeImage(2, "Calque", "b"),
        FakeImage(3, "Calque", "c"),
    ]
    check(deck.anchor_collisions(repeated) == ["Calque"],
          f"a repeated anchor is reported (got {deck.anchor_collisions(repeated)})")
    check(deck.anchor_collisions([FakeImage(1, "", "a"), FakeImage(2, "", "b")]) == [],
          "empty anchors are not reported as colliding with each other")


def test_assemble() -> None:
    print("\nassemble")
    images = [
        FakeImage(1, "Diapositive 1", "aaa"),
        FakeImage(2, "Calque", "logo"),
        FakeImage(3, "Calque", "logo"),
        FakeImage(4, "Diapositive 4", "logo"),
        FakeImage(5, "Sans numéro", "ccc"),
    ]
    manifest = deck.assemble(images, FakeCheck(), en_offset=2)

    check(manifest["image_count"] == 5 and manifest["slide_count"] == 2,
          f"chrome is set aside, not counted as reviewable "
          f"(got {manifest['slide_count']} of {manifest['image_count']})")
    check(manifest["chrome_count"] == 3, "chrome is counted, never silently dropped")
    check(all(slide["is_chrome"] for slide in manifest["slides"] if slide["digest"].startswith("logo")),
          "every chrome image is marked as such")

    first = manifest["slides"][0]
    check(first["slide"] == 1 and first["numbered_by_document"],
          "a numbered anchor keys the slide")
    check(first["en_page"] == 3, f"en_offset maps slide 1 to page 3 (got {first['en_page']})")

    last = manifest["slides"][-1]
    check(last["slide"] == 5 and not last["numbered_by_document"],
          "an unnumbered anchor falls back to image order")
    check(manifest["unnumbered"] == [5],
          f"the fallback is reported so it can be checked (got {manifest['unnumbered']})")
    check(manifest["anchor_collisions"] == ["Calque"],
          "collisions travel with the manifest")
    check(not any(slide["anchor_usable"] for slide in manifest["slides"]
                  if slide["anchor"] == "Calque"),
          "a colliding anchor is marked unusable for comments")

    no_offset = deck.assemble(images, FakeCheck())
    check(all(slide["en_page"] is None for slide in no_offset["slides"]),
          "without an offset no English page is invented")


def test_comment_building() -> None:
    print("\ncomment building")
    manifest = deck.assemble([
        FakeImage(1, "Diapositive 1", "a"),
        FakeImage(2, "Diapositive 2", "b"),
        FakeImage(3, "Calque", "c"),
        FakeImage(4, "Calque", "d"),
    ], FakeCheck())

    findings = [
        {"slide": 1, "type": "Terminology", "current": "santé sexuelle",
         "proposed": "santé et droits sexuels et reproductifs",
         "why": "Drops the rights half of the term."},
        {"slide": 1, "type": "Register", "current": "vous devez",
         "proposed": "vous pouvez", "why": "Directive where the English is not."},
        {"slide": 2, "type": "Data gap", "why": "No English source page for this slide."},
        {"slide": 3, "type": "Spelling", "current": "indicatuer", "proposed": "indicateur"},
        {"slide": 99, "type": "Spelling", "current": "x", "proposed": "y"},
    ]

    grouped = deck.merge_findings(findings)
    check(len(grouped[1]) == 2, "two findings on one slide group together")
    check(deck.merge_findings([{"slide": "not a number"}]) == {},
          "a finding with no usable slide key is dropped from grouping")

    requests, skipped = deck.build_requests(findings, manifest)
    check(len(requests) == 1,
          f"only the anchorable slide gets a comment (got {len(requests)})")
    check(requests[0].match == "Diapositive 1",
          "the comment matches the slide-title paragraph")
    text = requests[0].text
    check(text.count("\n") == 1, "both findings on slide 1 share one balloon")
    check("« santé sexuelle »" in text and "« vous pouvez »" in text,
          "renderings are quoted in guillemets")
    check("Drops the rights half" in text, "the reason travels into the comment")

    joined = " | ".join(skipped)
    check("slide 2" in joined and "worklist only" in joined,
          f"a data gap is routed to the sheet, not a comment (got {joined})")
    check("slide 3" in joined and "repeats" in joined,
          "a colliding anchor is reported rather than mis-anchored")
    check("slide 99" in joined and "not in the deck" in joined,
          "a finding for a slide that does not exist is reported, not guessed at")

    empty, notes = deck.build_requests([], manifest)
    check(empty == [] and notes == [], "no findings means no requests and no noise")


def test_worklist_rows() -> None:
    print("\nworklist rows")
    manifest = deck.assemble([FakeImage(1, "Diapositive 1", "a"),
                              FakeImage(2, "Diapositive 2", "b")],
                             FakeCheck(), en_offset=10)
    findings = [
        {"slide": 2, "type": "Data gap", "why": "No source page."},
        {"slide": 1, "type": "Terminology", "current": "a", "proposed": "b",
         "severity": "High", "en_source": "SRHR"},
        {"slide": 1, "type": "Spelling", "current": "c", "proposed": "d"},
    ]
    rows = deck.worklist_rows(findings, manifest)

    check(len(rows) == 3, f"the sheet keeps every finding, queries included (got {len(rows)})")
    check([row["Slide"] for row in rows] == ["1", "1", "2"], "rows come out in slide order")
    check(list(rows[0]) == deck.SHEET_COLUMNS,
          f"columns match the agreed schema (got {list(rows[0])})")
    check(rows[0]["EN page"] == "11", "the slide's English page carries into the row")
    check(rows[2]["Severity"] == "Medium",
          "a finding with no severity defaults rather than blanking")
    check(rows[0]["Severity"] == "High", "a stated severity is preserved")


# --------------------------------------------------------------------------

REL = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _fixture(tmp: Path) -> Path:
    """A Storyline-shaped export whose titles carry French and Romanian accents."""
    titles = ["Diapositive 1 : Chișinău", "Diapositive 2 : Résumé",
              "Diapositive 3 : Données", "Diapositive 4 : Références"]

    def paragraph(text: str = "", rid: str | None = None) -> str:
        inner = f"<w:r><w:t>{text}</w:t></w:r>" if text else ""
        if rid:
            inner += (
                '<w:r><w:drawing><wp:inline xmlns:wp="http://schemas.openxmlformats'
                '.org/drawingml/2006/wordprocessingDrawing">'
                '<a:graphic xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">'
                '<a:graphicData><pic:pic xmlns:pic="http://schemas.openxmlformats'
                '.org/drawingml/2006/picture"><pic:blipFill>'
                f'<a:blip xmlns:r="{REL}" r:embed="{rid}"/>'
                '</pic:blipFill></pic:pic></a:graphicData></a:graphic>'
                '</wp:inline></w:drawing></w:r>'
            )
        return f"<w:p>{inner}</w:p>"

    body = "".join(paragraph(title) + paragraph(rid=f"rId{n}")
                   for n, title in enumerate(titles, start=1))
    document = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        f'<w:document xmlns:w="{W}"><w:body>{body}</w:body></w:document>'
    )
    rels = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        + "".join(f'<Relationship Id="rId{n}" Type="{REL}/image" '
                  f'Target="../media/image{n}.bin"/>' for n in range(1, 5))
        + "</Relationships>"
    )
    root_rels = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        f'<Relationship Id="rIdMain" Type="{REL}/officeDocument" '
        'Target="word/document2.xml"/></Relationships>'
    )

    path = tmp / "export_fr.docx"
    with zipfile.ZipFile(path, "w") as archive:
        archive.writestr("[Content_Types].xml",
                         '<?xml version="1.0"?><Types xmlns="http://schemas.'
                         'openxmlformats.org/package/2006/content-types"/>')
        archive.writestr("_rels/.rels", root_rels)
        archive.writestr("word/document2.xml", document)
        archive.writestr("word/_rels/document2.xml.rels", rels)
        for n in range(1, 5):
            archive.writestr(f"media/image{n}.bin",
                             b"\xff\xd8\xff\xe0" + f"SLIDE{n}".encode() + b"\x00" * 32)
    return path


def test_cli_subprocess(tmp: Path) -> None:
    print("\nCLI as a subprocess — the console encoding is real")
    fixture = _fixture(tmp)
    out_dir = tmp / "slides"

    proc = subprocess.run(
        [sys.executable, str(Path(__file__).resolve().parent / "extract_deck.py"),
         "extract", str(fixture), "--out-dir", str(out_dir), "--en-offset", "4"],
        capture_output=True, text=True, encoding="utf-8", timeout=120,
    )
    check(proc.returncode == 0, f"extract exits 0 (stderr: {proc.stderr[:200]})")
    # The Wave 3 bug, exactly: an accented title through a cp1252 console.
    check("Chișinău" in proc.stdout,
          "an accented slide title survives the console encoding")
    check("Diapositive 4" in proc.stdout, "every slide reaches the table")
    check("| 1 |" in proc.stdout and "| 4 |" in proc.stdout,
          "slides are keyed by their own French labels")
    check(len(list(out_dir.glob("*.jpeg"))) == 4,
          f"images are written with the sniffed extension "
          f"(got {[p.name for p in out_dir.iterdir()]})")

    json_proc = subprocess.run(
        [sys.executable, str(Path(__file__).resolve().parent / "extract_deck.py"),
         "extract", str(fixture), "--out-dir", str(out_dir), "--json"],
        capture_output=True, text=True, encoding="utf-8", timeout=120,
    )
    manifest = json.loads(json_proc.stdout)
    check(manifest["slide_count"] == 4, "the JSON manifest carries every slide")
    check(manifest["is_deck_export"], "the fixture is recognised as a deck export")

    findings = [{"slide": 2, "type": "Terminology", "current": "santé sexuelle",
                 "proposed": "santé et droits sexuels et reproductifs",
                 "why": "Controlled glossary rendering.", "severity": "High"}]
    findings_path = tmp / "findings.json"
    findings_path.write_text(json.dumps(findings, ensure_ascii=False), encoding="utf-8")

    sheet_path = tmp / "corrections.xlsx"
    sheet_proc = subprocess.run(
        [sys.executable, str(Path(__file__).resolve().parent / "extract_deck.py"),
         "sheet", str(fixture), "--findings", str(findings_path),
         "--out", str(sheet_path), "--date", "04/08/2026"],
        capture_output=True, text=True, encoding="utf-8", timeout=120,
    )
    check(sheet_proc.returncode == 0 and sheet_path.is_file(),
          f"sheet writes the workbook (stderr: {sheet_proc.stderr[:200]})")
    if sheet_path.is_file():
        from openpyxl import load_workbook
        ws = load_workbook(sheet_path)["Corrections"]
        header = [cell.value for cell in ws[3]]
        check(header == deck.SHEET_COLUMNS, f"the sheet carries the agreed columns (got {header})")
        check(ws["F4"].value == "santé et droits sexuels et reproductifs",
              "accented French reaches the written cell intact")

    bad = subprocess.run(
        [sys.executable, str(Path(__file__).resolve().parent / "extract_deck.py"),
         "extract", str(tmp / "missing.docx"), "--out-dir", str(out_dir)],
        capture_output=True, text=True, encoding="utf-8", timeout=120,
    )
    check(bad.returncode == 2, f"a missing file exits 2 (got {bad.returncode})")


def _png(rgb: tuple[int, int, int]) -> bytes:
    """A valid 1x1 PNG, built here so the fixtures need no imaging library.

    Distinct colours matter: python-docx stores identical bytes once, and three
    identical slide images would be read back as repeated chrome.
    """
    import struct
    import zlib

    def chunk(tag: bytes, data: bytes) -> bytes:
        body = tag + data
        return (struct.pack(">I", len(data)) + body
                + struct.pack(">I", zlib.crc32(body) & 0xFFFFFFFF))

    ihdr = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
    return (b"\x89PNG\r\n\x1a\n" + chunk(b"IHDR", ihdr)
            + chunk(b"IDAT", zlib.compress(b"\x00" + bytes(rgb)))
            + chunk(b"IEND", b""))


def test_cli_comment(tmp: Path) -> None:
    print("\ncomment subcommand — the anchor has to land on a real paragraph")
    try:
        from docx import Document
    except ImportError:
        check(True, "comment checks skipped (python-docx not installed)")
        return

    titles = ["Diapositive 1 : Introduction", "Diapositive 2 : Santé",
              "Diapositive 3 : Références"]
    document = Document()
    for number, title in enumerate(titles, start=1):
        image = tmp / f"pixel{number}.png"
        image.write_bytes(_png((number * 60, 40, 200)))
        document.add_paragraph(title)
        document.add_picture(str(image))
    source = tmp / "deck_standard.docx"
    document.save(str(source))

    findings = [
        {"slide": 2, "type": "Terminology", "current": "santé sexuelle",
         "proposed": "santé et droits sexuels et reproductifs",
         "why": "Controlled glossary rendering.", "severity": "High"},
        {"slide": 3, "type": "Data gap", "why": "No English source page."},
    ]
    findings_path = tmp / "comment_findings.json"
    findings_path.write_text(json.dumps(findings, ensure_ascii=False), encoding="utf-8")

    commented = tmp / "deck_commented.docx"
    proc = subprocess.run(
        [sys.executable, str(Path(__file__).resolve().parent / "extract_deck.py"),
         "comment", str(source), "--findings", str(findings_path),
         "--out", str(commented)],
        capture_output=True, text=True, encoding="utf-8", timeout=120,
    )
    check(proc.returncode == 0 and commented.is_file(),
          f"comment writes a commented copy (stderr: {proc.stderr[:200]})")
    check(source.is_file() and "Diapositive" not in _comments_of(source),
          "the original is left without comments")
    check("slide 3" in proc.stderr and "worklist only" in proc.stderr,
          f"the data gap is reported as routed to the sheet (stderr: {proc.stderr[:200]})")

    if commented.is_file():
        text = _comments_of(commented)
        check("« santé et droits sexuels et reproductifs »" in text,
              f"the proposed rendering reaches the balloon (got {text[:160]!r})")
        check(text.count("Terminology") == 1,
              "one comment per slide, and the data gap did not become one")


def _comments_of(path: Path) -> str:
    from ane_package.officeops import read_comments

    try:
        return "\n".join(comment.text for comment in read_comments(path))
    except Exception:
        return ""


def main() -> int:
    tmp = Path(tempfile.mkdtemp(prefix="deck-mode-tests-"))
    test_slide_numbering()
    test_chrome_and_collisions()
    test_assemble()
    test_comment_building()
    test_worklist_rows()
    test_cli_subprocess(tmp)
    test_cli_comment(tmp)

    total = PASSED + FAILED
    print(f"\n{PASSED}/{total} checks passed (localise deck mode)")
    print(f"artefacts: {tmp}")
    return 1 if FAILED else 0


if __name__ == "__main__":
    raise SystemExit(main())
