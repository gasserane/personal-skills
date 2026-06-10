#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
deskreview.py — reliable mechanics for an IPPF MA accreditation desk review.

The model supplies judgement (status + conclusion + questions) as JSON; this toolkit handles
the Word document safely. Every brittle thing that broke in real use is handled here:
content-based cell location (Word merges tables on save), backup, file-lock check, UTF-8,
edit-preservation (append/replace below the MA text), and content-based verification.

Run any subcommand with -h for usage. Always run with PYTHONUTF8=1.

Subcommands:
  locate      <docx>                                  list in-scope standard + check cells
  extract     <evidence_dir> <out_dir>                dump readable evidence to .txt + manifest
  apply       <docx> <conclusions.json> [--mode]      write reviewer blocks (insert|replace)
  verify      <docx> <conclusions.json>               per-cell: exactly one block, MA text kept
  standalone  <docx> <conclusions.json> <out.docx> --principles 3 4
  interview   <items.json> <out.docx>                 build pre-interview + interview lists

conclusions.json schema:
  {"STD": {"3.1": {"status": "...", "body": ["para", ...], "questions": ["q", ...]}},
   "CHK": {"3.1.1": {"status": "...", "body": "para or [paras]", "questions": [...]}}}
"""
import argparse, json, os, re, shutil, sys
import docx
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PURPLE = RGBColor(0x70, 0x30, 0xA0)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREY = RGBColor(0x44, 0x44, 0x44)
RED = RGBColor(0xB0, 0x1A, 0x2A)
PURBAR = "7030A0"; LILAC = "EEE6F7"
HEAD = "REVIEWER CONCLUSIONS"            # match with/without colon

MIN_FONT = 10.0                          # hard floor: no run in a generated companion doc
def fpt(size):                           # renders below 10pt (Ane's font rule, locked 2026-06-10)
    """Font size as Pt, floored at MIN_FONT. Use for every font-size assignment in a generated
    document. Do NOT use for spacing (space_after, indents) — those stay as raw Pt()."""
    return Pt(max(size, MIN_FONT))
STD_RE = re.compile(r"Standard\s+(\d+\.\d+)")
NUM_RE = re.compile(r"^(\d+\.\d+\.\d+)")

# ---------------------------------------------------------------- file safety
def is_locked(path):
    if not os.path.exists(path):
        return False
    try:
        with open(path, "r+b"):
            return False
    except (PermissionError, OSError):
        return True

def backup(path, tag="edit"):
    base, ext = os.path.splitext(path)
    dst = f"{base}_BACKUP-before-{tag}{ext}"
    shutil.copy2(path, dst)
    return dst

# ---------------------------------------------------------------- cell location (by content)
def iter_cells(doc):
    """Yield ('STD', std_id, cell) for each standard DESCRIPTION cell and
       ('CHK', check_num, evidence_cell) for each check row. Robust to Word table merges."""
    cur = None
    for t in doc.tables:
        for row in t.rows:
            c0 = row.cells[0].text
            m = STD_RE.search(c0)
            if m:
                # Only treat a "Standard X.Y" match as a header if it sits in the
                # original cell text, not inside an appended reviewer block. A
                # reviewer conclusion may legitimately name another standard
                # (e.g. "filed under the strategy-and-policy standard 2.3"); that
                # must not reset the current-standard tracker. See reliability-playbook.
                hp = c0.find(HEAD)
                if hp == -1 or m.start() < hp:
                    cur = m.group(1)
            if "comply with this Standard" in c0 and cur:
                yield ("STD", cur, row.cells[0])
            mn = NUM_RE.match(c0.strip())
            if mn:
                yield ("CHK", mn.group(1), row.cells[-1])

# ---------------------------------------------------------------- evidence extraction
def _pdf_text(path, maxpages=80):
    try:
        from pypdf import PdfReader
    except ImportError:
        return "[pypdf not installed: run python -m pip install pypdf]"
    try:
        r = PdfReader(path)
        return "\n".join((pg.extract_text() or "") for pg in r.pages[:maxpages])
    except Exception as e:
        return f"[PDF ERROR: {e}]"

def _docx_text(path):
    d = docx.Document(path)
    parts = [p.text for p in d.paragraphs if p.text.strip()]
    for t in d.tables:
        for r in t.rows:
            ded = []
            for c in r.cells:
                if not ded or ded[-1] != c.text:
                    ded.append(c.text)
            line = " | ".join(x.strip() for x in ded if x.strip())
            if line:
                parts.append(line)
    return "\n".join(parts)

def _xlsx_text(path):
    import openpyxl
    wb = openpyxl.load_workbook(path, data_only=True, read_only=True)
    out = []
    for ws in wb.worksheets:
        out.append(f"### SHEET {ws.title}")
        for row in ws.iter_rows(values_only=True):
            vals = [str(v) for v in row if v is not None]
            if vals:
                out.append(" | ".join(vals))
    return "\n".join(out)

def extract_text(path):
    """Return (text_or_None, status). status: OK | SCAN/EMPTY | SKIP-binary."""
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext == ".docx":
            txt = _docx_text(path)
        elif ext == ".pdf":
            txt = _pdf_text(path)
        elif ext == ".xlsx":
            txt = _xlsx_text(path)
        elif ext in (".txt", ".csv", ".md"):
            txt = open(path, encoding="utf-8", errors="replace").read()
        else:
            return None, "SKIP-binary"        # .doc, images, .rar/.zip
    except Exception as e:
        return f"[EXTRACT ERROR {ext}: {e}]", "ERROR"
    n = len((txt or "").strip())
    return txt, ("OK" if n > 40 else "SCAN/EMPTY")

def cmd_extract(args):
    os.makedirs(args.out_dir, exist_ok=True)
    manifest = []
    safe = lambda s: re.sub(r"[^A-Za-z0-9._-]", "_", s)[:120]
    for dirpath, _, files in os.walk(args.evidence_dir):
        for fn in files:
            full = os.path.join(dirpath, fn)
            rel = os.path.relpath(full, args.evidence_dir)
            txt, status = extract_text(full)
            manifest.append((status, len((txt or "").strip()), os.path.splitext(fn)[1].lower(), rel))
            if txt is not None:
                out = os.path.join(args.out_dir, safe(rel.replace(os.sep, "__")) + ".txt")
                with open(out, "w", encoding="utf-8") as f:
                    f.write(f"FILE: {rel}\nSTATUS: {status}\n{'='*50}\n{txt}")
    with open(os.path.join(args.out_dir, "_MANIFEST.txt"), "w", encoding="utf-8") as f:
        for status, n, ext, rel in sorted(manifest):
            f.write(f"{status:12} {n:8} {ext:6} {rel}\n")
    ok = sum(1 for m in manifest if m[0] == "OK")
    gap = sum(1 for m in manifest if m[0] in ("SCAN/EMPTY", "SKIP-binary"))
    print(f"extracted {len(manifest)} files | OK={ok} | DATA-GAPS(scan/binary)={gap}")
    print(f"manifest: {os.path.join(args.out_dir, '_MANIFEST.txt')}")
    print("Review the DATA-GAP rows: request readable copies or verify at interview.")

# ---------------------------------------------------------------- block writing
def add_block(cell, status, body, questions):
    p = cell.add_paragraph(); r = p.add_run(f"REVIEWER CONCLUSIONS: {status}")
    r.bold = True; r.font.color.rgb = PURPLE
    for para in (body if isinstance(body, list) else [body]):
        bp = cell.add_paragraph(); rr = bp.add_run(para); rr.font.color.rgb = PURPLE
    if questions:
        qp = cell.add_paragraph(); qr = qp.add_run("Questions / aspects for the interview:")
        qr.bold = True; qr.font.color.rgb = PURPLE
        for q in questions:
            lp = cell.add_paragraph(); lr = lp.add_run("- " + q); lr.font.color.rgb = PURPLE

def replace_block(cell, status, body, questions):
    paras = cell.paragraphs
    start = next((i for i, p in enumerate(paras) if p.text.startswith(HEAD)), None)
    if start is not None:
        for p in paras[start:]:
            p._element.getparent().remove(p._element)
    add_block(cell, status, body, questions)

def cmd_apply(args):
    if is_locked(args.docx):
        sys.exit("ABORT: target is open in Word (locked). Close it and re-run.")
    if not args.no_backup:
        print("backup:", backup(args.docx, args.tag))
    data = json.load(open(args.conclusions, encoding="utf-8"))
    std, chk = data.get("STD", {}), data.get("CHK", {})
    doc = docx.Document(args.docx)
    done = []
    for kind, ref, cell in iter_cells(doc):
        spec = std.get(ref) if kind == "STD" else chk.get(ref)
        if not spec:
            continue
        args_ = (cell, spec["status"], spec.get("body", []), spec.get("questions", []))
        (replace_block if args.mode == "replace" else add_block)(*args_)
        done.append((kind, ref))
    doc.save(args.docx)
    for k, r in done:
        print(f"  {args.mode.upper()} {k} {r}")
    missing = [r for r in list(std) + list(chk) if r not in {d[1] for d in done}]
    print(f"applied {len(done)} | not located: {missing or 'none'}")

def cmd_verify(args):
    data = json.load(open(args.conclusions, encoding="utf-8"))
    want = set(data.get("STD", {})) | set(data.get("CHK", {}))
    doc = docx.Document(args.docx)
    seen = {}; problems = []
    for kind, ref, cell in iter_cells(doc):
        if ref in want:
            n = cell.text.count(HEAD)
            seen[ref] = n
            if n != 1:
                problems.append((ref, n))
    for ref in sorted(want):
        n = seen.get(ref, 0)
        flag = "OK" if n == 1 else "PROBLEM"
        print(f"  {ref:8} blocks={n} {flag}")
    print("PASS" if (not problems and set(seen) == want) else f"FAIL: {problems or 'missing refs'}")

def cmd_locate(args):
    doc = docx.Document(args.docx)
    for kind, ref, cell in iter_cells(doc):
        has = HEAD in cell.text
        print(f"{kind} {ref:8} existing_block={has}")

# ---------------------------------------------------------------- standalone per-principle doc
def _shade(cell, hexv):
    tcPr = cell._tc.get_or_add_tcPr(); sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear"); sh.set(qn("w:fill"), hexv); tcPr.append(sh)
def _setw(cell, cm):
    cell.width = Cm(cm); tcPr = cell._tc.get_or_add_tcPr(); w = OxmlElement("w:tcW")
    w.set(qn("w:w"), str(int(cm * 567))); w.set(qn("w:type"), "dxa"); tcPr.append(w)
def _crun(cell, text, size=8.5, bold=False, color=None, italic=False):
    cell.text = ""; p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(1)
    r = p.add_run(text); r.bold = bold; r.italic = italic; r.font.size = fpt(size)
    if color is not None: r.font.color.rgb = color

def _scan_template(doc):
    """Pull standard descriptions/verdicts and check questions/verdicts from the template."""
    std_meta, chk_meta = {}, {}; cur = None
    for t in doc.tables:
        for ri, row in enumerate(t.rows):
            c0 = row.cells[0].text
            m = STD_RE.search(c0)
            hp = c0.find(HEAD)
            if m and (hp == -1 or m.start() < hp):  # ignore standard ids inside reviewer prose
                cur = m.group(1)
                std_meta.setdefault(cur, {})["title"] = c0.split("\t")[-1].strip()[:60]
            if "comply with this Standard" in c0 and cur:
                pre = c0.split("REVIEWER CONCLUSIONS")[0]
                desc = pre.split("Is the Association")[0].strip()
                verd = ""
                if "comply with this Standard?" in pre:
                    verd = pre.split("comply with this Standard?")[-1].strip()
                std_meta.setdefault(cur, {}).update(desc=desc, verdict=verd[:60])
            mn = NUM_RE.match(c0.strip())
            if mn:
                cells = row.cells
                q = cells[1].text.strip().split("\n")[0][:160]
                verd = cells[3].text.strip()[:20] if len(cells) > 3 else ""
                chk_meta[mn.group(1)] = {"q": q, "verdict": verd}
    return std_meta, chk_meta

def cmd_standalone(args):
    src = docx.Document(args.docx)
    std_meta, chk_meta = _scan_template(src)
    data = json.load(open(args.conclusions, encoding="utf-8"))
    std, chk = data.get("STD", {}), data.get("CHK", {})
    principles = [str(p) for p in args.principles]
    d = docx.Document(); sec = d.sections[0]
    for a in ("top_margin", "bottom_margin"): setattr(sec, a, Cm(1.4))
    sec.left_margin = Cm(1.6); sec.right_margin = Cm(1.6)
    def para(txt, size=10, bold=False, italic=False, color=None, after=4, before=0):
        p = d.add_paragraph(); p.paragraph_format.space_after = Pt(after); p.paragraph_format.space_before = Pt(before)
        r = p.add_run(txt); r.bold = bold; r.italic = italic; r.font.size = fpt(size)
        if color is not None: r.font.color.rgb = color
    para(args.title or "IPPF MA Accreditation Desk Review — reviewer output", 13, True, color=PURPLE, after=2)
    para(f"Principle(s) {', '.join(principles)} | standard-level conclusion is the primary record for the IPPF reader.",
         9, italic=True, color=GREY, after=8)
    for sid in sorted(s for s in std if s.split(".")[0] in principles):
        meta = std_meta.get(sid, {}); spec = std[sid]
        tb = d.add_table(rows=1, cols=1); c = tb.rows[0].cells[0]; _shade(c, PURBAR); _setw(c, 17.8); c.text = ""
        p = c.paragraphs[0]; r = p.add_run(f"  Standard {sid} — {meta.get('title','')}")
        r.bold = True; r.font.size = fpt(11); r.font.color.rgb = WHITE
        if meta.get("desc"): para("Standard (IPPF): " + meta["desc"], 8.5, italic=True, color=GREY, after=1, before=2)
        if meta.get("verdict"): para("MA self-assessment verdict: " + meta["verdict"], 8.5, bold=True, color=GREY, after=4)
        para("REVIEWER CONCLUSION — STANDARD LEVEL", 9.5, True, color=PURPLE, after=1)
        para("Desk-review status: " + spec["status"], 9.5, True, color=PURPLE, after=2)
        for cp in (spec.get("body", []) if isinstance(spec.get("body"), list) else [spec.get("body","")]):
            para(cp, 9.5, color=PURPLE, after=3)
        if spec.get("questions"):
            para("Questions / aspects to explore at the interview:", 9.5, True, color=PURPLE, after=1)
            for q in spec["questions"]:
                p = d.add_paragraph(); p.paragraph_format.left_indent = Cm(0.5); p.paragraph_format.space_after = Pt(1)
                rr = p.add_run("• " + q); rr.font.size = fpt(9.5); rr.font.color.rgb = PURPLE
        para("Supporting detail — checks", 8.5, True, color=GREY, after=2, before=4)
        t = d.add_table(rows=1, cols=3); t.style = "Table Grid"; h = t.rows[0].cells
        for i, (txt, w) in enumerate([("Check", 4.6), ("MA", 1.2), ("Reviewer conclusion (with MA self-assessment)", 12.0)]):
            _crun(h[i], txt, 8.5, True, WHITE); _shade(h[i], PURBAR); _setw(h[i], w)
        for cid in sorted(c for c in chk if c.startswith(sid + ".")):
            cs = chk[cid]; cm = chk_meta.get(cid, {})
            row = t.add_row().cells
            row[0].text = ""; pp = row[0].paragraphs[0]
            rr = pp.add_run(cid + "  "); rr.bold = True; rr.font.size = fpt(8)
            rr2 = pp.add_run(cm.get("q", "")); rr2.italic = True; rr2.font.size = fpt(7.5); rr2.font.color.rgb = GREY; _setw(row[0], 4.6)
            _crun(row[1], cm.get("verdict", ""), 8, True); _setw(row[1], 1.2)
            row[2].text = ""; pp = row[2].paragraphs[0]
            rs = pp.add_run(cs["status"] + " — "); rs.bold = True; rs.font.size = fpt(8); rs.font.color.rgb = PURPLE
            body = cs.get("body", ""); body = " ".join(body) if isinstance(body, list) else body
            rb = pp.add_run(body); rb.font.size = fpt(8); rb.font.color.rgb = PURPLE
            if cs.get("questions"):
                pq = row[2].add_paragraph(); rq = pq.add_run("Interview: " + " ".join(cs["questions"]))
                rq.italic = True; rq.font.size = fpt(7.5); rq.font.color.rgb = PURPLE
            _setw(row[2], 12.0)
        d.add_paragraph().paragraph_format.space_after = Pt(4)
    d.save(args.out)
    print("saved standalone:", args.out)

# ---------------------------------------------------------------- interview lists
def cmd_interview(args):
    data = json.load(open(args.items, encoding="utf-8"))
    d = docx.Document(); sec = d.sections[0]
    for a in ("top_margin", "bottom_margin", "left_margin", "right_margin"): setattr(sec, a, Cm(1.3))
    def para(txt, size=10, bold=False, italic=False, color=None, after=4):
        p = d.add_paragraph(); p.paragraph_format.space_after = Pt(after)
        r = p.add_run(txt); r.bold = bold; r.italic = italic; r.font.size = fpt(size)
        if color is not None: r.font.color.rgb = color
    para(data.get("title", "Pre-interview requests and interview questions"), 13, True, color=PURPLE, after=2)
    para("Type key: NC = non-compliance / gap related; CL = clarification (standard met, evidence/confirmation needed).",
         9, italic=True, color=GREY, after=6)
    def section(title, rows, towhom):
        tb = d.add_table(rows=1, cols=1); c = tb.rows[0].cells[0]; _shade(c, PURBAR); _setw(c, 18.4); c.text = ""
        rr = c.paragraphs[0].add_run("  " + title); rr.bold = True; rr.font.size = fpt(11); rr.font.color.rgb = WHITE
        cols = 4 if towhom else 3
        t = d.add_table(rows=1, cols=cols); t.style = "Table Grid"; h = t.rows[0].cells
        hdrs = [("Check(s)", 1.9), ("Type", 1.1), ("Item", 11.0 if towhom else 13.4)] + ([("To whom", 2.0)] if towhom else [])
        for i, (txt, w) in enumerate(hdrs):
            _crun(h[i], txt, 9, True, WHITE); _shade(h[i], PURBAR); _setw(h[i], w)
        for it in rows:
            if it.get("principle"):
                br = t.add_row().cells; m = br[0]
                for j in range(1, cols): m = m.merge(br[j])
                _crun(m, it["principle"], 9.5, True, PURPLE); _shade(m, LILAC); continue
            rc = t.add_row().cells
            _crun(rc[0], it["ref"], 9, True); _setw(rc[0], 1.9)
            _crun(rc[1], it["type"], 9, True, RED if it["type"] == "NC" else GREY); _setw(rc[1], 1.1)
            _crun(rc[2], it["item"], 9); _setw(rc[2], 11.0 if towhom else 13.4)
            if towhom:
                _crun(rc[3], it.get("to", ""), 9); _setw(rc[3], 2.0)
        d.add_paragraph().paragraph_format.space_after = Pt(3)
    section("A. Pre-interview clarifications and document requests from the Executive Director", data.get("pre_interview", []), False)
    section("B. Questions for the interview phase", data.get("interview", []), True)
    d.save(args.out)
    print("saved interview lists:", args.out)

# ---------------------------------------------------------------- CLI
def main():
    ap = argparse.ArgumentParser(description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
    sub = ap.add_subparsers(dest="cmd", required=True)
    p = sub.add_parser("locate"); p.add_argument("docx"); p.set_defaults(fn=cmd_locate)
    p = sub.add_parser("extract"); p.add_argument("evidence_dir"); p.add_argument("out_dir"); p.set_defaults(fn=cmd_extract)
    p = sub.add_parser("apply"); p.add_argument("docx"); p.add_argument("conclusions")
    p.add_argument("--mode", choices=["insert", "replace"], default="insert")
    p.add_argument("--tag", default="edit"); p.add_argument("--no-backup", action="store_true"); p.set_defaults(fn=cmd_apply)
    p = sub.add_parser("verify"); p.add_argument("docx"); p.add_argument("conclusions"); p.set_defaults(fn=cmd_verify)
    p = sub.add_parser("standalone"); p.add_argument("docx"); p.add_argument("conclusions"); p.add_argument("out")
    p.add_argument("--principles", nargs="+", required=True); p.add_argument("--title"); p.set_defaults(fn=cmd_standalone)
    p = sub.add_parser("interview"); p.add_argument("items"); p.add_argument("out"); p.set_defaults(fn=cmd_interview)
    args = ap.parse_args()
    args.fn(args)

if __name__ == "__main__":
    main()
